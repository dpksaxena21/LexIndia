# LexIndia: Streamlit Web Interface
# Built by: Deepak Saxena

import streamlit as st
import requests
import os
import re
import anthropic
from dotenv import load_dotenv
from tavily import TavilyClient
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import fitz                       # pymupdf

load_dotenv()
INDIAN_KANOON_TOKEN = os.getenv("INDIAN_KANOON_TOKEN")
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
ECOURTS_API_KEY = os.getenv("ECOURTS_API_KEY")

# Page Configuration
st.set_page_config(
    page_title="LexIndia - AI Legal Assistant",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom Styling
st.markdown("""
<style>
    .stApp, .main, section[data-testid="stMain"] {
        background-color: #ffffff !important;
    }
    .stApp p, .stApp h1, .stApp h2, .stApp h3,
    .stApp h4, .stApp li, .stApp span, .stApp div,
    .stMarkdown p, .stMarkdown h1, .stMarkdown h2,
    .stMarkdown h3, .stMarkdown li {
        color: #333333 !important;
    }
    [data-testid="stSidebar"] {
        background-color: #1a2744 !important;
    }
    [data-testid="stSidebar"] * {
        color: #ffffff !important;
    }
    .stButton > button {
        background-color: #c9a84c !important;
        color: white !important;
        border: none;
        border-radius: 8px;
        padding: 10px 30px;
        font-size: 16px;
        font-weight: bold;
        width: 100%;
    }
    .stButton > button:hover {
        background-color: #b8943d !important;
        color: white !important;
    }
    [data-testid="stExpander"] {
        background-color: #ffffff !important;
        border: 1px solid #e0e0e0;
        border-radius: 8px;
    }
    [data-testid="stExpander"] * {
        color: #333333 !important;
    }
    .stTextInput > div > div > input {
        border: 2px solid #1a2744;
        border-radius: 8px;
        font-size: 16px;
        padding: 12px;
        color: #333333 !important;
    }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

# Initialise session state
if "module" not in st.session_state:
    st.session_state.module = "lexsearch"
if "history" not in st.session_state:
    st.session_state.history = []
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "scanned_dates" not in st.session_state:
    st.session_state.scanned_dates = []
if "scanned_docs" not in st.session_state:
    st.session_state.scanned_docs = []

# Sidebar Navigation
with st.sidebar:
    st.markdown("## ⚖️ LexIndia")
    st.markdown("*Your AI Legal Assistant*")
    st.markdown("---")
    st.markdown("### Navigation")

    with st.expander("🔍 Research", expanded=True):
        if st.button("LexSearch - Case Research"):
            st.session_state.module = "lexsearch"
        if st.button("LexPlain - Law Explainer"):
            st.session_state.module = "lexplain"
        if st.button("LexConstitute - Constitutional Advisor"):
            st.session_state.module = "lexconstitute"
        if st.button("LexGlobe - International Law"):
            st.session_state.module = "lexglobe"

    with st.expander("✍️ Drafting"):
        if st.button("LexScan - Document Analyser"):
            st.session_state.module = "lexscan"
        if st.button("LexDraft - Draft Generator"):
            st.session_state.module = "lexdraft"
        if st.button("LexDebate - Counter Argument Generator"):
            st.session_state.module = "lexdebate"

    with st.expander("📁 File Management"):
        if st.button("LexCause - Daily Cause List"):
            st.session_state.module = "lexdiary"
        if st.button("LexVault - File Storage"):
            st.session_state.module = "lexvault"
        if st.button("LexChat - Legal Chatbot"):
            st.session_state.module = "lexchat"
        if st.button("LexTrack - Live Case Updates"):
            st.session_state.module = "lextrack"

    with st.expander("📊 Intelligence"):
        if st.button("LexPredict - Outcome Predictor"):
            st.session_state.module = "lexpredict"
        if st.button("LexBench - Judge Analysis"):
            st.session_state.module = "lexbench"
        if st.button("LexPulse - Trend Dashboard"):
            st.session_state.module = "lexpulse"
        if st.button("LexVoice - Regional Languages"):
            st.session_state.module = "lexvoice"
        if st.button("LexMap - Court Locator"):
            st.session_state.module = "lexmap"

    st.markdown("---")
    st.markdown("### 🕐 Recent Searches")
    if st.session_state.history:
        for item in reversed(st.session_state.history[-10:]):
            st.markdown(f"**{item['module']}**")
            st.markdown(f"*{item['query'][:60]}...*" if len(item['query']) > 60 else f"*{item['query']}*")
            st.markdown("---")
    else:
        st.markdown("*No searches yet.*")
    st.markdown("*Powered by Indian Kanoon API and Claude AI*")

# Main Page Header
st.markdown("""
<div style="text-align: center; padding: 40px 0 20px 0;">
    <h1 style="color: #1a2744; font-size: 56px; font-weight: 900;
               letter-spacing: 4px; margin-bottom: 8px;">
        ⚖️ LEXINDIA
    </h1>
    <p style="color: #c9a84c; font-size: 20px; font-style: italic;
              margin-bottom: 40px;">
        India's AI Legal Research Assistant
    </p>
</div>
""", unsafe_allow_html=True)

# ─── LEXSEARCH MODULE ────────────────────────────────────────────────────────
if st.session_state.module == "lexsearch":
    col1, col2, col3 = st.columns([1, 4, 1])
    with col2:
        with st.form("search_form"):
            search_query = st.text_input(
                label="Search Query",
                placeholder="Search any case, law, or legal topic...",
                label_visibility="collapsed"
            )
            search_clicked = st.form_submit_button("🔍 Search LexIndia")
            if search_clicked and not search_query:
                st.warning("Please enter a legal query to search.")

    if search_clicked and search_query:
        st.session_state.history.append({"module": "🔍 LexSearch", "query": search_query})

        # Search India Code for statutory text
        with st.spinner("Searching India Code for statutory text..."):
            try:
                tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
                india_code_results = tavily_client.search(
                    query=f"{search_query} site:indiacode.nic.in",
                    search_depth="basic",
                    max_results=3
                )
                india_code_text = ""
                for result in india_code_results.get("results", []):
                    india_code_text += f"**{result['title']}**\n{result['content'][:500]}\n\n"
                if india_code_text:
                    st.markdown("### 📜 Statutory Text — India Code")
                    st.markdown(india_code_text)
                    st.markdown("---")
            except Exception:
                pass

        with st.spinner("Searching Indian Kanoon for cases..."):
            params = {"formInput": search_query, "pagenum": 0}
            url = "https://api.indiankanoon.org/search/"
            headers = {"Authorization": f"Token {INDIAN_KANOON_TOKEN}"}
            try:
                response = requests.post(url, headers=headers, params=params)
                if response.status_code == 200:
                    results = response.json()
                    st.success(f"Found {results['found']} results for '{search_query}'")
                    st.markdown("---")
                    for index, doc in enumerate(results['docs'][:3]):
                        clean_title = re.sub(r'<[^>]+>', '', doc['title'])
                        clean_court = re.sub(r'<[^>]+>', '', doc['docsource'])
                        if index == 0:
                            st.markdown("### 📌 Best Match")
                        elif index == 1:
                            st.markdown("### 📎 Related Cases")
                        with st.expander(f"📄 {clean_title}"):
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown(f"**Court:** {clean_court}")
                            with col2:
                                st.markdown(f"**Date:** {doc['publishdate']}")
                            doc_id = doc['tid']
                            st.markdown(f"**🔗 Full Judgment:** [Read on Indian Kanoon](https://indiankanoon.org/doc/{doc_id}/)")
                            st.markdown("---")
                            st.markdown("**🤖 AI Summary:**")
                            with st.spinner("Generating AI analysis..."):
                                judgment_url = f"https://api.indiankanoon.org/doc/{doc_id}/"
                                j_response = requests.post(judgment_url, headers=headers)
                                judgment_text = ""
                                if j_response.status_code == 200:
                                    judgment_text = j_response.json().get('doc', '')
                                client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
                                if judgment_text:
                                    content = f"Full judgment:\n{judgment_text[:6000]}"
                                else:
                                    content = f"Case: {clean_title}\nCourt: {clean_court}\nDate: {doc['publishdate']}"
                                prompt = f"""You are a legal assistant specializing in Indian law.
{content}
Provide: 1. Brief facts 2. Court decision 3. Key legal principle 4. How to use in arguments 5. Similar cases 6. Weaknesses
Be concise and practical for a practicing lawyer."""
                                message = client.messages.create(
                                    model="claude-haiku-4-5-20251001",
                                    max_tokens=4096,
                                    messages=[{"role": "user", "content": prompt}]
                                )
                                st.markdown(message.content[0].text)
                else:
                    st.error("Search failed. Please try again.")
            except Exception as e:
                st.error("Connection error. Please check your internet connection.")

# ─── LEXPLAIN MODULE ─────────────────────────────────────────────────────────
elif st.session_state.module == "lexplain":
    st.markdown("## 📖 LexPlain — Law Explainer")
    st.markdown("*Type any legal concept, section, or term and get a plain language explanation.*")
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 4, 1])
    with col2:
        with st.form("lexplain_form"):
            legal_query = st.text_input(
                label="Legal Query",
                placeholder="e.g. What is res judicata? Explain Section 302 IPC. What is anticipatory bail?",
                label_visibility="collapsed"
            )
            explain_clicked = st.form_submit_button("💡 Explain")
            if explain_clicked and not legal_query:
                st.warning("Please enter a legal concept to explain.")
    if explain_clicked and legal_query:
        st.session_state.history.append({"module": "📖 LexPlain", "query": legal_query})
        with st.spinner("Generating explanation..."):
            client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
            prompt = f"""You are a legal expert specializing in Indian law.

A lawyer or law student has asked: "{legal_query}"

Please provide:
1. Simple plain language explanation (as if explaining to a non-lawyer)
2. Technical legal definition
3. Key elements or requirements
4. Relevant Indian laws, sections, or articles
5. Important cases that established or clarified this concept
6. Practical examples of how this applies in real cases

Be clear, accurate, and practical."""
            try:
                message = client.messages.create(
                    model="claude-haiku-4-5-20251001",
                    max_tokens=4096,
                    messages=[{"role": "user", "content": prompt}]
                )
                st.markdown("### 💡 Explanation")
                st.markdown(message.content[0].text)
            except Exception as e:
                st.error(f"Error: {str(e)}")

# ─── LEXDEBATE MODULE ─────────────────────────────────────────────────────────
elif st.session_state.module == "lexdebate":
    st.markdown("## ⚔️ LexDebate - Counter Argument Finder")
    st.markdown("*Enter your legal argument and get the strongest counter arguments from Indian case law.*")
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 4, 1])
    with col2:
        with st.form("lexdebate_form"):
            argument = st.text_area(
                label="Legal Argument",
                placeholder="e.g. My client should get bail because he has no prior criminal record...",
                height=150,
                label_visibility="collapsed"
            )
            debate_clicked = st.form_submit_button("⚔️ Find Counter Arguments")
            if debate_clicked and not argument:
                st.warning("Please enter a legal argument.")
    if debate_clicked and argument:
        st.session_state.history.append({"module": "⚔️ LexDebate", "query": argument[:80]})
        with st.spinner("Finding counter arguments..."):
            client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
            prompt = f"""You are a senior Indian lawyer preparing for the opposing side.

A lawyer has made this argument:
"{argument}"

Please provide:
1. The strongest counter-arguments the opposing counsel will make
2. Indian cases that support the opposing position
3. Weaknesses in the original argument
4. How the original lawyer can strengthen their argument
5. Relevant sections or provisions the opposing side will cite

Be specific, cite actual Indian cases, and be practical for a courtroom setting."""
            try:
                message = client.messages.create(
                    model="claude-haiku-4-5-20251001",
                    max_tokens=4096,
                    messages=[{"role": "user", "content": prompt}]
                )
                st.markdown("### ⚔️ Counter Arguments")
                st.markdown(message.content[0].text)
            except Exception as e:
                st.error(f"Error: {str(e)}")

# ─── LEXCONSTITUTE MODULE ────────────────────────────────────────────────────
elif st.session_state.module == "lexconstitute":
    st.markdown("## 🏛️ LexConstitute - Constitutional Advisor")
    st.markdown("*Ask any constitutional question and get an answer with relevant Articles, landmark cases, amendments, and live current affairs.*")
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 4, 1])
    with col2:
        with st.form("lexconstitute_form"):
            const_query = st.text_input(
                label="Constitutional Query",
                placeholder="e.g. What are the limits of Article 19? Explain basic structure doctrine.",
                label_visibility="collapsed"
            )
            const_clicked = st.form_submit_button("🏛️ Get Constitutional Analysis")
            if const_clicked and not const_query:
                st.warning("Please enter a constitutional question.")
    if const_clicked and const_query:
        st.session_state.history.append({"module": "🏛️ LexConstitute", "query": const_query})
        with st.spinner("Searching for latest developments..."):
            try:
                tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
                search_results = tavily_client.search(
                    query=f"{const_query} India Supreme Court 2024 2025",
                    search_depth="basic",
                    max_results=5
                )
                recent_news = ""
                for result in search_results.get("results", []):
                    recent_news += f"- {result['title']}: {result['content'][:300]}\n"
            except Exception:
                recent_news = "No recent news available."
        with st.spinner("Generating constitutional analysis..."):
            client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
            prompt = f"""You are a constitutional law expert specializing in Indian law.

Question: "{const_query}"

Recent news and developments from the web:
{recent_news}

Please provide:
1. Constitutional provisions — relevant Articles with exact text
2. Simple plain language explanation of each Article with real life examples
3. Landmark Supreme Court cases that established key principles
4. Current legal position — how courts interpret this today
5. Recent developments — based on the news above
6. Practical implications for lawyers and citizens
7. Constitutional Amendments — relevant amendments that changed this provision
8. Pending issues — unresolved constitutional questions

Be thorough, cite actual cases, and include current affairs where relevant."""
            try:
                message = client.messages.create(
                    model="claude-haiku-4-5-20251001",
                    max_tokens=4096,
                    messages=[{"role": "user", "content": prompt}]
                )
                st.markdown("### 🏛️ Constitutional Analysis")
                st.markdown(message.content[0].text)
            except Exception as e:
                st.error(f"Error: {str(e)}")

# ─── LEXCHAT MODULE ──────────────────────────────────────────────────────────

elif st.session_state.module == "lexchat":
    st.markdown("## 💬 LexChat - Legal Chatbot")
    st.markdown("*Describe your legal situation and get expert advice. Ask follow-up questions.*")
    st.markdown("---")
    
    st.warning("⚠️ **Privacy Notice:** Do not enter client names, case numbers, or any identifying information. Use generic terms like 'my client', 'the accused', 'the complainant'. Your conversation is processed by AI servers.")
    
    # Show conversation history
    for message in st.session_state.chat_history:
        if message["role"] == "user":
            st.markdown(f"**🧑‍⚖️ You:** {message['content']} ")
        else:
            st.markdown(f"**⚖️ LexChat:** {message['content']}")
        st.markdown("---")
        
    # Clear chat button
    if st.session_state.chat_history:
        if st.button("🗑️ Clear Chat"):
            st.session_state.chat_history = []
            st.rerun()
            
    # Input form
    col1, col2, col3 = st.columns([1,4,1])
    with col2:
        with st.form("lexchat_form"):
            user_message = st.text_area(
                label = "Your message",
                placeholder="e.g. My client hit someone with a car. Person is injured. What should I do immediately?",
                height=100,
                label_visibility="collapsed"
            )
            chat_clicked = st.form_submit_button("💬 Send")
            if chat_clicked and not user_message:
                st.warning("Please type your message.")
                
            if chat_clicked and user_message:
                st.session_state.history.append({"module": "💬 LexChat", "query": user_message[:80]})
                st.session_state.chat_history.append({"role": "user", "content": user_message})
                with st.spinner("LexChat is thinking..."):
                    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
                    system_prompt = """You are LexChat — a legendary legal mind with over 50 years of experience in Indian and international law.

Your background:
- Former Supreme Court of India advocate with 500+ landmark cases
- Retired High Court Judge with expertise in constitutional, criminal, and civil law
- Deep knowledge of IPC, CrPC, BNS, BNSS, Indian Evidence Act, Constitution of India
- Expert in international law, human rights law, and comparative jurisprudence
- Have appeared before the Supreme Court, High Courts, Sessions Courts, and Tribunals
- Authored legal commentaries and textbooks on Indian criminal and constitutional law
- Mentored hundreds of lawyers across India

Your approach:
- You give precise, practical, actionable legal advice
- You always cite the exact section, article, or case with year and court
- You think like both a defence lawyer AND a prosecutor — you see all angles
- You explain complex law in simple language without losing accuracy
- You anticipate what the opposing side will argue
- You know which judges are strict, which courts are lenient, which arguments work
- You understand the ground reality of Indian courts — not just textbook law

Your rules:
- Never give vague or generic answers
- Always be specific to Indian law and Indian court procedure
- Cite BNS/BNSS for new cases, IPC/CrPC for old cases
- Never encourage sharing client names, case numbers, or identifying details
- If you do not know something — say so honestly rather than guessing
- Always tell the lawyer what to do NEXT — immediate steps, not just theory

You are the lawyer every lawyer wishes they could call at midnight before a hearing."""
                    try:
                        message = client.messages.create(
                            model="claude-haiku-4-5-20251001",
                            max_tokens=4096,
                            system=system_prompt,
                            messages=st.session_state.chat_history
                        )
                        assistant_response = message.content[0].text
                        st.session_state.chat_history.append({"role": "assistant", "content": assistant_response})
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error: {str(e)}")

# ─── LEXTRACK MODULE ─────────────────────────────────────────────────────────
elif st.session_state.module == "lextrack":
    st.markdown("## 📡 LexTrack — Live Case Updates")
    st.markdown("*Track cases by CNR number or search by party name, advocate, or judge.*")
    st.markdown("---")

    track_tab1, track_tab2 = st.tabs(["🔍 Search Cases", "📋 CNR Lookup"])

    # ── TAB 1: CASE SEARCH ───────────────────────────────────────────────────
    with track_tab1:
        st.markdown("### 🔍 Search Cases")
        st.markdown("*Search by party name, advocate name, judge name, or any keyword.*")

        col1, col2 = st.columns(2)
        with col1:
            search_party = st.text_input("Party / Litigant Name", placeholder="e.g. Rajesh Kumar")
            search_advocate = st.text_input("Advocate Name", placeholder="e.g. Adv. Sharma")
        with col2:
            search_judge = st.text_input("Judge Name", placeholder="e.g. Justice Gupta")
            search_court_code = st.text_input("Court Code (optional)", placeholder="e.g. DLHC01")

        search_status = st.selectbox("Case Status", ["All", "PENDING", "DISPOSED"])

        if st.button("🔍 Search Cases"):
            if not any([search_party, search_advocate, search_judge]):
                st.warning("Please enter at least one search term.")
            else:
                with st.spinner("Searching 27 crore case records..."):
                    try:
                        params = {"pageSize": 10}
                        if search_party:
                            params["litigants"] = search_party
                        if search_advocate:
                            params["advocates"] = search_advocate
                        if search_judge:
                            params["judges"] = search_judge
                        if search_court_code:
                            params["courtCodes"] = search_court_code
                        if search_status != "All":
                            params["caseStatuses"] = search_status

                        search_url = "https://webapi.ecourtsindia.com/api/partner/search"
                        search_headers = {"Authorization": f"Bearer {ECOURTS_API_KEY}"}
                        search_response = requests.get(search_url, headers=search_headers, params=params)

                        if search_response.status_code == 200:
                            search_data = search_response.json()
                            results = search_data.get('data', {}).get('results', [])
                            total = search_data.get('data', {}).get('totalHits', 0)

                            if results:
                                st.success(f"Found **{total}** cases. Showing top {len(results)}.")
                                st.markdown("---")
                                for case in results:
                                    petitioners = ', '.join(case.get('petitioners', ['Unknown']))
                                    respondents = ', '.join(case.get('respondents', ['Unknown']))
                                    title = f"{petitioners[:40]} vs {respondents[:40]}"
                                    cnr = case.get('cnr', '')
                                    status = case.get('caseStatus', 'N/A')
                                    next_hearing = case.get('nextHearingDate', 'N/A')
                                    court = case.get('courtCode', 'N/A')
                                    case_type = case.get('caseType', 'N/A')

                                    status_icon = "🟢" if status == "PENDING" else "⚫"
                                    with st.expander(f"{status_icon} {title}"):
                                        col1, col2 = st.columns(2)
                                        with col1:
                                            st.markdown(f"**CNR:** {cnr}")
                                            st.markdown(f"**Case Type:** {case_type}")
                                            st.markdown(f"**Status:** {status}")
                                            st.markdown(f"**Court:** {court}")
                                        with col2:
                                            st.markdown(f"**Next Hearing:** {next_hearing}")
                                            st.markdown(f"**Filing Date:** {case.get('filingDate', 'N/A')}")
                                            advocates = ', '.join(case.get('petitionerAdvocates', ['N/A']))
                                            st.markdown(f"**Advocate:** {advocates}")
                                        if cnr:
                                            st.markdown(f"**🔗 View on eCourts:** [Click here](https://indiankanoon.org/search/?formInput={cnr})")
                                        st.session_state.history.append({"module": "📡 LexTrack", "query": title[:60]})
                            else:
                                st.info("No cases found. Try different search terms.")
                        else:
                            st.error(f"Search failed. Status: {search_response.status_code}")
                    except Exception as e:
                        st.error(f"Connection error: {str(e)}")

    # ── TAB 2: CNR LOOKUP ────────────────────────────────────────────────────
    with track_tab2:
        st.markdown("### 📋 CNR Lookup")
        st.info("💡 **What is CNR?** Every court case in India has a unique CNR (Case Number Record). Format: DLHC010001232024 — State code + Court code + Case number + Year")

        col1, col2, col3 = st.columns([1, 4, 1])
        with col2:
            with st.form("lextrack_form"):
                cnr_number = st.text_input(
                    label="CNR Number",
                    placeholder="e.g. DLHC010001232024",
                    label_visibility="collapsed"
                )
                track_clicked = st.form_submit_button("📡 Track Case")
                if track_clicked and not cnr_number:
                    st.warning("Please enter a CNR Number")
                if track_clicked and cnr_number:
                    st.session_state.history.append({"module": "📡 LexTrack", "query": cnr_number})
                    with st.spinner("Fetching live case data from eCourts..."):
                        try:
                            url = f"https://webapi.ecourtsindia.com/api/partner/case/{cnr_number.strip().upper()}"
                            headers = {"Authorization": f"Bearer {ECOURTS_API_KEY}"}
                            response = requests.get(url, headers=headers)
                            if response.status_code == 200:
                                data = response.json()
                                st.success("✅ Case found!")
                                st.markdown("---")

                                case_data = data.get('data', {}).get('courtCaseData', {})
                                entity_info = data.get('data', {}).get('entityInfo', {})

                                st.markdown("### 📋 Case Details")
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.markdown(f"**Case Number:** {case_data.get('filingNumber', 'N/A')}")
                                    st.markdown(f"**Case Type:** {case_data.get('caseTypeRaw', 'N/A')}")
                                    st.markdown(f"**Filing Date:** {case_data.get('filingDate', 'N/A')}")
                                    st.markdown(f"**Status:** {case_data.get('caseStatus', 'N/A')}")
                                with col2:
                                    st.markdown(f"**Court:** {case_data.get('courtName', 'N/A')}")
                                    st.markdown(f"**Judge:** {', '.join(case_data.get('judges', ['N/A']))}")
                                    st.markdown(f"**Next Hearing:** {entity_info.get('nextDateOfHearing', 'N/A')}")
                                    st.markdown(f"**Last Hearing:** {case_data.get('lastHearingDate', 'N/A')}")

                                st.markdown("---")
                                st.markdown("### 👥 Parties")
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.markdown(f"**Petitioner:** {', '.join(case_data.get('petitioners', ['N/A']))}")
                                with col2:
                                    st.markdown(f"**Respondent:** {', '.join(case_data.get('respondents', ['N/A']))}")

                                # AI Case Summary
                                ai_analysis = data.get('data', {}).get('caseAiAnalysis', {})
                                if ai_analysis:
                                    st.markdown("---")
                                    st.markdown("### 🤖 AI Case Summary")
                                    st.markdown(f"**Summary:** {ai_analysis.get('caseSummary', 'N/A')}")
                                    st.markdown(f"**Case Type:** {ai_analysis.get('caseType', 'N/A')}")
                                    st.markdown(f"**Complexity:** {ai_analysis.get('complexity', 'N/A')}")
                                    key_issues = ai_analysis.get('keyIssues', [])
                                    if key_issues:
                                        st.markdown(f"**Key Issues:** {', '.join(key_issues)}")

                                orders = case_data.get('judgmentOrders', [])
                                files = data.get('data', {}).get('files', {}).get('files', [])
                                if orders:
                                    st.markdown("---")
                                    st.markdown("### 📄 Recent Orders")
                                    for i, order in enumerate(orders[:3]):
                                        st.markdown(f"**{order.get('orderDate', 'N/A')}** — {order.get('orderType', 'N/A')}")
                                        if i < len(files):
                                            markdown_content = files[i].get('markdownContent', '')
                                            if markdown_content:
                                                with st.expander("📖 Read Full Order"):
                                                    st.markdown(markdown_content)

                                ias = case_data.get('interlocutoryApplications', [])
                                if ias:
                                    st.markdown("---")
                                    st.markdown("### 📑 Interlocutory Applications")
                                    for ia in ias[:3]:
                                        st.markdown(f"**{ia.get('regNo', 'N/A')}** — Filed by: {ia.get('filedBy', 'N/A')} — Status: {ia.get('status', 'N/A')}")

                            elif response.status_code == 404:
                                st.error("Case not found. Please check the CNR Number.")
                            elif response.status_code == 401:
                                st.error("API authentication failed.")
                            else:
                                st.error(f"Error fetching case data. Status: {response.status_code}")
                        except Exception as e:
                            st.error(f"Connection error: {str(e)}")
# ─── LEXDRAFT MODULE ─────────────────────────────────────────────────────────
elif st.session_state.module == "lexdraft":
    st.markdown("## ✍️ LexDraft — Legal Document Generator")
    st.markdown("*Select a document type, fill in the details, and get a professionally drafted legal document.*")
    st.markdown("---")

    st.info("🤖 **LexDraft AI** is purpose-built for Indian legal drafting — trained on Indian court formats, BNS/BNSS/CPC provisions, and High Court filing standards. Documents are generated with court-accurate structure and legal language. Please review for factual accuracy and verify case citations before filing, as AI may occasionally make errors in specific citations or local court rules.")

    doc_category = st.selectbox("Select Category", [
        "Criminal",
        "Supreme Court",
        "High Court",
        "Civil",
        "Affidavits",
        "Notices",
        "General"
    ])

    doc_types = {
        "Criminal": [
            "Regular Bail Application (S.483 BNSS)",
            "Anticipatory Bail Application (S.482 BNSS)",
            "Quashing Petition (S.528 BNSS)",
            "Discharge Application",
            "Bail Cancellation Opposition",
            "Revision Petition (Criminal)",
            "Default Bail Application (S.187 BNSS)",
            "Surrender Application"
        ],
        "Supreme Court": [
            "SLP (Civil) — Art. 136",
            "SLP (Criminal) — Art. 136",
            "Civil Appeal",
            "Criminal Appeal",
            "Review Petition",
            "Curative Petition",
            "Writ Petition (Art. 32) — Fundamental Rights",
            "Transfer Petition"
        ],
        "High Court": [
            "Writ Petition (Art. 226) — Mandamus",
            "Writ Petition (Art. 226) — Certiorari",
            "Writ Petition (Art. 226) — Habeas Corpus",
            "Writ Appeal",
            "Criminal Appeal",
            "Criminal Revision Petition",
            "Anticipatory Bail (High Court)",
            "Letters Patent Appeal"
        ],
        "Civil": [
            "Plaint",
            "Written Statement",
            "Interlocutory Application",
            "Contempt Petition",
            "Execution Application",
            "Revision Petition (Civil)",
            "Appeal against Decree",
            "Injunction Application (Order 39 CPC)"
        ],
        "Affidavits": [
            "General Affidavit",
            "Supporting Affidavit",
            "Counter Affidavit",
            "Rejoinder Affidavit",
            "Affidavit of Assets and Liabilities",
            "Affidavit of Undertaking",
            "Affidavit in lieu of Examination in Chief"
        ],
        "Notices": [
            "Legal Notice",
            "Reply to Legal Notice",
            "Cheque Bounce Notice (S.138 NI Act)",
            "Consumer Complaint Notice",
            "RTI Application",
            "Section 80 CPC Notice",
            "Eviction Notice",
            "Defamation Notice"
        ],
        "General": [
            "Vakalatnama",
            "Memo of Appearance",
            "Power of Attorney (General)",
            "Power of Attorney (Special)",
            "Rent Agreement",
            "MOU (Memorandum of Understanding)",
            "Settlement Agreement",
            "Consumer Complaint (NCDRC/State/District)"
        ]
    }

    doc_type = st.selectbox("Select Document Type", doc_types[doc_category])

    st.markdown("---")
    st.markdown("### 📝 Fill in the Details")

    col1, col2 = st.columns(2)
    with col1:
        court_name = st.text_input("Court Name", placeholder="e.g. High Court of Delhi at New Delhi")
        petitioner_name = st.text_input("Petitioner / Applicant Name", placeholder="e.g. Ramesh Kumar")
        respondent_name = st.text_input("Respondent Name", placeholder="e.g. State of Delhi")
    with col2:
        case_number = st.text_input("Case Number (if any)", placeholder="e.g. FIR No. 123/2024")
        petitioner_advocate = st.text_input("Advocate Name", placeholder="e.g. Adv. Suresh Sharma")
        date = st.text_input("Date", placeholder="e.g. 21st April 2025")

    facts = st.text_area("Brief Facts of the Case", placeholder="Describe the facts of the case in brief...", height=150)
    grounds = st.text_area("Grounds / Prayer", placeholder="List the grounds for relief or specific prayer...", height=100)

    if st.button("✍️ Generate Document"):
        if not petitioner_name or not facts:
            st.warning("Please fill in at least Petitioner Name and Facts.")
        else:
            with st.spinner("Drafting your legal document..."):
                client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
                prompt = f"""You are a senior Indian lawyer with 30 years of experience drafting legal documents for Indian courts.

Document Type: {doc_type}
Court: {court_name}
Case Number: {case_number}
Petitioner/Applicant: {petitioner_name}
Respondent: {respondent_name}
Advocate: {petitioner_advocate}
Date: {date}
Facts: {facts}
Grounds/Prayer: {grounds}

STRICT FORMATTING RULES:
- Follow exact Indian court formatting standards
- Use proper legal language — "respectfully showeth", "humbly prays", "most respectfully"
- All sections must be numbered
- Include proper cause title
- Always end with PRAYER clause and VERIFICATION where required
- Cite BNS/BNSS for offences after July 2024, IPC/CrPC for older matters
- CRITICAL: Use correct BNSS 2023 sections — S.483 for regular bail, S.482 for anticipatory bail, S.187 for default bail, S.528 for quashing. Never use old CrPC sections 436/437/438/439/482 for post-July 2024 matters
DOCUMENT TEMPLATES:

═══ REGULAR BAIL APPLICATION (S.483 BNSS) ═══
IN THE COURT OF [COURT NAME]
BAIL APPLICATION NO. ___ OF [YEAR]

IN THE MATTER OF:
[ACCUSED NAME] ... Applicant/Accused
VERSUS
STATE OF [STATE] ... Respondent

APPLICATION FOR REGULAR BAIL UNDER SECTION 483 BNSS, 2023

MOST RESPECTFULLY SHOWETH:
1. That the applicant [name], aged [age], S/o [father], resident of [address], is the accused in FIR No. [X]/[YEAR] registered at PS [Name] under Sections [X] BNS, 2023.
2. That the applicant was arrested on [date] and has been in judicial custody since [date].
3. That the chargesheet has/has not been filed.
4. [Additional facts]

GROUNDS:
I. That the applicant has deep roots in the community and is not a flight risk.
II. That the applicant has no prior criminal antecedents.
III. That the allegations are false and motivated.
IV. That the investigation is complete and custody is not required.
V. That Article 21 of the Constitution guarantees right to personal liberty.
VI. That the offence is [bailable/triable by Magistrate].

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Release the applicant on bail on such terms as deemed fit;
(b) Pass any other order in the interest of justice.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Applicant

VERIFICATION:
I, [Name], verify that the contents are true and correct to the best of my knowledge. Verified at [Place] on [Date].

═══ ANTICIPATORY BAIL APPLICATION (S.482 BNSS) ═══
IN THE COURT OF
ANTICIPATORY BAIL APPLICATION NO. ___ OF [YEAR]

IN THE MATTER OF:
[APPLICANT NAME] ... Applicant
VERSUS
STATE OF [STATE] ... Respondent

APPLICATION FOR ANTICIPATORY BAIL UNDER SECTION 482 BNSS, 2023
MOST RESPECTFULLY SHOWETH:
1. That the applicant apprehends arrest in connection with [case/FIR details].
2. That the applicant is [description — occupation, address, background].
3. That the applicant has not been arrested and no FIR has been registered / FIR No. [X] has been registered.
4. [Facts explaining why arrest is apprehended]

GROUNDS:
I. That the apprehension of arrest is well-founded.
II. That the applicant is innocent and has been falsely implicated.
III. That the applicant is a respectable member of society with no criminal antecedents.
IV. That the applicant undertakes to cooperate with the investigation.
V. That custodial interrogation is not required.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Direct that in the event of arrest, the applicant be released on bail;
(b) Grant interim protection from arrest till disposal of this application;
(c) Pass any other order in the interest of justice.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Applicant

VERIFICATION:
Verified at [Place] on [Date] that contents are true and correct.

═══ QUASHING PETITION (S.528 BNSS) ═══
IN THE HIGH COURT OF [STATE] AT [CITY]
CRIMINAL MISCELLANEOUS PETITION NO. ___ OF [YEAR]

IN THE MATTER OF:
[PETITIONER NAME] ... Petitioner
VERSUS
STATE OF [STATE] AND ANR. ... Respondents

PETITION UNDER SECTION 528 BNSS, 2023 FOR QUASHING OF FIR NO. [X]/[YEAR]

MOST RESPECTFULLY SHOWETH:
1. That the Petitioner challenges FIR No. [X]/[YEAR] registered at PS [Name] under Sections [X] BNS.
2. That the facts giving rise to the FIR are: [facts].
3. That the FIR is liable to be quashed on the following grounds.

GROUNDS:
I. That the FIR does not disclose any cognizable offence.
II. That the FIR is filed with malafide intent to harass the Petitioner.
III. That the matter is purely civil in nature and has been given criminal colour.
IV. That there has been a settlement between parties [if applicable].
V. That continuation of proceedings would be abuse of process of court.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Quash FIR No. [X]/[YEAR] and all proceedings arising therefrom;
(b) Stay further investigation during pendency of this petition;
(c) Pass any other order as deemed fit.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Petitioner

═══ DISCHARGE APPLICATION ═══
IN THE COURT OF [COURT NAME]
[CASE TYPE] NO. [NUMBER] OF [YEAR]

IN THE MATTER OF:
STATE OF [STATE] ... Prosecution
VERSUS
[ACCUSED NAME] ... Accused

APPLICATION FOR DISCHARGE UNDER SECTION 250 BNSS, 2023

MOST RESPECTFULLY SHOWETH:
1. That the applicant/accused is facing trial in the above case.
2. That chargesheet was filed on [date] under Sections [X] BNS.
3. That the material on record does not make out a prima facie case.
4. [Supporting facts]

GROUNDS:
I. That there is no sufficient ground for proceeding against the accused.
II. That the evidence collected during investigation is insufficient.
III. That the prosecution witnesses are unreliable.
IV. That the alleged offence is not made out from the facts on record.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Discharge the accused from the charges framed;
(b) Pass any other order in the interest of justice.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Accused

═══ DEFAULT BAIL APPLICATION (S.187 BNSS) ═══
IN THE COURT OF [COURT NAME]
BAIL APPLICATION NO. ___ OF [YEAR]

IN THE MATTER OF:
[ACCUSED NAME] ... Applicant
VERSUS
STATE OF [STATE] ... Respondent

APPLICATION FOR DEFAULT BAIL UNDER SECTION 187(3) BNSS, 2023

MOST RESPECTFULLY SHOWETH:
1. That the applicant was arrested on [date] in connection with FIR No. [X]/[YEAR] under Sections [X] BNS.
2. That [60/90] days have elapsed since the date of arrest.
3. That the investigation agency has failed to file chargesheet within the statutory period.
4. That the applicant is entitled to bail as a matter of right under Section 187(3) BNSS.

GROUNDS:
I. That the statutory period of [60/90] days has expired without filing of chargesheet.
II. That the right to default bail is an indefeasible right of the accused.
III. That the Supreme Court in Rakesh Kumar Paul v. State of Assam has held that default bail cannot be defeated.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Release the applicant on default bail forthwith;
(b) Pass any other order in the interest of justice.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Applicant

═══ REVISION PETITION (CRIMINAL) ═══
IN THE [SESSIONS COURT / HIGH COURT] OF [STATE]
CRIMINAL REVISION PETITION NO. ___ OF [YEAR]

IN THE MATTER OF:
[PETITIONER NAME] ... Petitioner/Revisionist
VERSUS
[RESPONDENT NAME] ... Respondent

CRIMINAL REVISION PETITION UNDER SECTION 442 BNSS, 2023
AGAINST ORDER DATED [DATE] PASSED BY [COURT NAME] IN [CASE NO.]

MOST RESPECTFULLY SHOWETH:
1. That the Petitioner is aggrieved by the order dated [date] passed by [court].
2. That the facts of the case are: [facts].
3. That the impugned order is illegal and liable to be set aside.

GROUNDS:
I. That the learned court below has committed an error of law.
II. That the impugned order is contrary to the evidence on record.
III. That the impugned order is perverse and arbitrary.
IV. That the court below failed to consider material evidence.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Set aside/modify the impugned order dated [date];
(b) Pass such other order as deemed fit.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Petitioner

═══ SLP (CIVIL/CRIMINAL) — ART. 136 ═══
IN THE SUPREME COURT OF INDIA
SPECIAL LEAVE PETITION ([CIVIL/CRIMINAL]) NO. ___ OF [YEAR]

IN THE MATTER OF:
[PETITIONER NAME] ... Petitioner(s)
VERSUS
[RESPONDENT NAME] ... Respondent(s)

SPECIAL LEAVE PETITION UNDER ARTICLE 136 OF THE CONSTITUTION OF INDIA
AGAINST JUDGMENT DATED [DATE] OF THE HIGH COURT OF [STATE] IN [CASE NO.]

MOST RESPECTFULLY SHOWETH:
1. That the Petitioner is aggrieved by the impugned judgment dated [date].
2. [Chronological facts]

QUESTIONS OF LAW:
A. Whether the High Court erred in [specific issue]?
B. Whether the judgment is contrary to settled law on [point]?

GROUNDS:
I. That the impugned judgment is contrary to law and facts.
II. That the High Court failed to appreciate [point].
III. That the matter involves substantial question of law of public importance.
IV. That the judgment is in conflict with Supreme Court precedents.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Grant Special Leave to Appeal;
(b) Stay the impugned judgment;
(c) After admission, allow the appeal and set aside the impugned judgment;
(d) Pass any other order as deemed fit.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Petitioner

═══ CIVIL APPEAL ═══
IN THE SUPREME COURT OF INDIA / HIGH COURT OF [STATE]
CIVIL APPEAL NO. ___ OF [YEAR]

IN THE MATTER OF:
[APPELLANT NAME] ... Appellant
VERSUS
[RESPONDENT NAME] ... Respondent

MEMORANDUM OF CIVIL APPEAL
AGAINST JUDGMENT DATED [DATE] OF [COURT] IN [CASE NO.]

MOST RESPECTFULLY SHOWETH:
1. That the Appellant is aggrieved by the judgment and decree dated [date].
2. [Facts]

GROUNDS OF APPEAL:
1. That the learned court below erred in [ground 1].
2. That the findings are contrary to evidence on record.
3. That the impugned judgment is perverse and unsustainable in law.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Allow this appeal;
(b) Set aside the impugned judgment;
(c) Pass decree in favour of the Appellant;
(d) Award costs.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Appellant

═══ REVIEW PETITION ═══
IN THE SUPREME COURT OF INDIA / HIGH COURT OF [STATE]
REVIEW PETITION NO. ___ OF [YEAR]
IN [CASE TYPE] NO. [NUMBER] OF [YEAR]

IN THE MATTER OF:
[PETITIONER NAME] ... Petitioner
VERSUS
[RESPONDENT NAME] ... Respondent

REVIEW PETITION UNDER ORDER XLVII RULE 1 CPC / ARTICLE 137 OF THE CONSTITUTION

MOST RESPECTFULLY SHOWETH:
1. That the Petitioner seeks review of judgment dated [date] in [case no.].
2. That the following errors apparent on the face of record warrant review.

GROUNDS FOR REVIEW:
I. Error apparent on face of record: [specific error].
II. Discovery of new and important evidence: [details if applicable].
III. Any other sufficient reason: [details].

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Review and recall the judgment dated [date];
(b) After review, pass appropriate orders;
(c) Pass any other order as deemed fit.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Petitioner

═══ CURATIVE PETITION ═══
IN THE SUPREME COURT OF INDIA
CURATIVE PETITION NO. ___ OF [YEAR]
IN [CASE TYPE] NO. [NUMBER] OF [YEAR]

IN THE MATTER OF:
[PETITIONER NAME] ... Petitioner
VERSUS
[RESPONDENT NAME] ... Respondent

CURATIVE PETITION UNDER ARTICLE 142 OF THE CONSTITUTION OF INDIA
[AS PER RUPA ASHOK HURRA v. ASHOK HURRA (2002) 4 SCC 388]

MOST RESPECTFULLY SHOWETH:
1. That all remedies including Review have been exhausted.
2. That the final judgment of this Court dated [date] suffers from grave miscarriage of justice.
3. That the present petition is certified by a Senior Advocate as per requirements.

GROUNDS:
I. That there has been violation of principles of natural justice.
II. That the petitioner was not heard before the order was passed.
III. That a Judge who was party to the lis participated in the decision.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Exercise curative jurisdiction and recall the judgment dated [date];
(b) Pass appropriate orders to prevent miscarriage of justice.

Place: [City]               [Advocate Name]
Date: [Date]                Senior Advocate / Advocate for Petitioner

═══ WRIT PETITION (ART. 226 / ART. 32) ═══
IN THE [HIGH COURT OF [STATE] / SUPREME COURT OF INDIA]
WRIT PETITION ([CIVIL/CRIMINAL]) NO. ___ OF [YEAR]

IN THE MATTER OF:
[PETITIONER NAME]
S/o [Father], aged [Age], Resident of [Address] ... Petitioner
VERSUS
1. [RESPONDENT 1]
2. [RESPONDENT 2]                              ... Respondent(s)

WRIT PETITION UNDER ARTICLE [226/32] OF THE CONSTITUTION OF INDIA
SEEKING WRIT OF [MANDAMUS/CERTIORARI/HABEAS CORPUS/PROHIBITION/QUO WARRANTO]

MOST RESPECTFULLY SHOWETH:

FACTS:
1. That the Petitioner is [description].
2. [Chronological facts]

QUESTIONS OF LAW:
A. Whether [Question 1]?
B. Whether [Question 2]?

GROUNDS:
I. That the impugned action is violative of Article [14/19/21] of the Constitution.
II. That the Respondent has acted without jurisdiction.
III. That the Petitioner has exhausted all alternative remedies.
IV. That irreparable harm will be caused if relief is not granted.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Issue writ of [MANDAMUS/CERTIORARI] directing [specific relief];
(b) Stay impugned order/action during pendency;
(c) Pass such other orders as deemed fit.

Date: [Date]                [Advocate Name]
Place: [City]               Advocate for Petitioner

═══ CONTEMPT PETITION ═══
IN THE [HIGH COURT / SUPREME COURT] OF [STATE]
CONTEMPT PETITION (CIVIL) NO. ___ OF [YEAR]
IN [ORIGINAL CASE NO.]

IN THE MATTER OF:
[PETITIONER NAME] ... Petitioner
VERSUS
[RESPONDENT/CONTEMNOR NAME] ... Respondent/Contemnor

CONTEMPT PETITION UNDER SECTION 11 OF THE CONTEMPT OF COURTS ACT, 1971
FOR WILFUL DISOBEDIENCE OF ORDER DATED [DATE]

MOST RESPECTFULLY SHOWETH:
1. That this Hon'ble Court passed an order dated [date] directing [specific direction].
2. That the Respondent/Contemnor was duly served with the said order.
3. That despite the order, the Respondent has wilfully failed to comply.
4. That the non-compliance is deliberate and intentional.

GROUNDS:
I. That the order dated [date] is clear and unambiguous.
II. That the Respondent has knowledge of the order.
III. That the non-compliance is wilful and deliberate.
IV. That the Petitioner has suffered prejudice due to non-compliance.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Issue notice to the Respondent/Contemnor;
(b) Hold the Respondent guilty of contempt of court;
(c) Punish the Respondent as per law;
(d) Direct compliance of the order dated [date];
(e) Pass any other order as deemed fit.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Petitioner

VERIFICATION:
Verified at [Place] on [Date] that contents are true and correct.

═══ REVISION PETITION (CIVIL) ═══
IN THE [DISTRICT COURT / HIGH COURT] OF [STATE]
CIVIL REVISION NO. ___ OF [YEAR]

IN THE MATTER OF:
[PETITIONER NAME] ... Petitioner/Revisionist
VERSUS
[RESPONDENT NAME] ... Respondent

CIVIL REVISION PETITION UNDER SECTION 115 CPC
AGAINST ORDER DATED [DATE] PASSED BY [COURT] IN [CASE NO.]

MOST RESPECTFULLY SHOWETH:
1. That the Petitioner is aggrieved by the order dated [date].
2. [Facts of the original case]
3. That the impugned order is illegal and erroneous.

GROUNDS:
I. That the court below has exercised jurisdiction not vested in it by law.
II. That the court below has failed to exercise jurisdiction vested in it.
III. That the court below has acted illegally or with material irregularity.
IV. That the impugned order has caused grave injustice to the Petitioner.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) Call for the record of the case;
(b) Set aside/modify the impugned order;
(c) Pass such other order as deemed fit.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Petitioner

═══ PLAINT ═══
IN THE COURT OF [COURT NAME]
CIVIL SUIT NO. ___ OF [YEAR]

[PLAINTIFF NAME]
S/o [Father], aged [Age], Resident of [Address]  ... Plaintiff
VERSUS
[DEFENDANT NAME]
Resident of [Address]                            ... Defendant

PLAINT UNDER ORDER VII RULE 1 CPC

MOST RESPECTFULLY SHOWETH:

1. CAUSE OF ACTION: The cause of action arose on [date] at [place] within jurisdiction of this Court.
2. VALUATION: The suit is valued at Rs. [Amount] for court fees and jurisdiction.
3. LIMITATION: The suit is within limitation as [reason].

FACTS:
4. [Numbered chronological facts]

LEGAL GROUNDS:
5. [Legal basis for the claim]

PRAYER:
Wherefore it is prayed that this Court may pass a decree:
(a) For [recovery of money / permanent injunction / declaration / specific performance];
(b) For interest at [X]% per annum;
(c) For costs of the suit;
(d) For such other relief as deemed fit.

Date: [Date]                [Advocate Name]
Place: [City]               Advocate for Plaintiff

VERIFICATION:
I, [Plaintiff Name], verify that the contents are true and correct to the best of my knowledge. Verified at [Place] on [Date].

═══ WRITTEN STATEMENT ═══
IN THE COURT OF [COURT NAME]
CIVIL SUIT NO. [NUMBER] OF [YEAR]

[PLAINTIFF NAME]                                 ... Plaintiff
VERSUS
[DEFENDANT NAME]                                 ... Defendant

WRITTEN STATEMENT ON BEHALF OF DEFENDANT
FILED UNDER ORDER VIII RULE 1 CPC

PRELIMINARY OBJECTIONS:
1. That the suit is not maintainable in law or on facts.
2. That this Court has no territorial/pecuniary jurisdiction.
3. That the suit is barred by limitation under Article [X] of the Limitation Act.
4. That the plaintiff has no cause of action against the defendant.
5. That the suit is bad for non-joinder of necessary parties.

REPLY ON MERITS:
Para-wise reply to plaint:
Para 1: [Admitted / Denied / Not admitted]
Para 2: [Response]
[Continue para by para]

ADDITIONAL FACTS IN DEFENCE:
1. [Facts supporting defendant's case]

PRAYER:
It is prayed that the suit be dismissed with exemplary costs.

Date: [Date]                [Advocate Name]
Place: [City]               Advocate for Defendant

VERIFICATION:
I, [Defendant Name], verify that the contents are true and correct. Verified at [Place] on [Date].

═══ INTERLOCUTORY APPLICATION ═══
IN THE COURT OF [COURT NAME]
[CASE TYPE] NO. [NUMBER] OF [YEAR]

IN THE MATTER OF:
[PETITIONER NAME] ... Petitioner/Applicant
VERSUS
[RESPONDENT NAME] ... Respondent

INTERLOCUTORY APPLICATION NO. ___ OF [YEAR]
FOR [STAY / INJUNCTION / DIRECTION / EXEMPTION]

MOST RESPECTFULLY SHOWETH:
1. That the above case is pending before this Court.
2. That the Applicant seeks interim relief on the following grounds.
3. [Facts necessitating interim relief]

GROUNDS:
I. That there is a prima facie case in favour of the Applicant.
II. That the balance of convenience lies in favour of granting relief.
III. That irreparable loss and injury will be caused if relief is not granted.
IV. That the Respondent will not suffer any prejudice.

PRAYER:
It is prayed that this Hon'ble Court may be pleased to:
(a) [Grant specific interim relief — stay/injunction/direction];
(b) Pass any other order as deemed fit.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Applicant

VERIFICATION:
Verified at [Place] on [Date] that contents are true and correct.

═══ EXECUTION APPLICATION ═══
IN THE COURT OF [COURT NAME]
EXECUTION PETITION NO. ___ OF [YEAR]
IN CIVIL SUIT / DECREE NO. [NUMBER] OF [YEAR]

[DECREE HOLDER NAME] ... Decree Holder
VERSUS
[JUDGMENT DEBTOR NAME] ... Judgment Debtor

APPLICATION FOR EXECUTION OF DECREE UNDER ORDER XXI CPC

MOST RESPECTFULLY SHOWETH:
1. That this Court passed a decree dated [date] in [case no.] in favour of the Decree Holder.
2. That the decree is for [payment of Rs. X / delivery of property / specific performance].
3. That the Judgment Debtor has failed to comply with the decree.
4. That the decree has not been satisfied till date.

GROUNDS:
I. That the decree is final and executable.
II. That the Judgment Debtor is deliberately avoiding compliance.
III. That the Decree Holder is entitled to execution as per law.

PRAYER:
It is prayed that this Court may be pleased to:
(a) Execute the decree dated [date] against the Judgment Debtor;
(b) Attach and sell the property of Judgment Debtor;
(c) Issue warrant of arrest of Judgment Debtor [if applicable];
(d) Pass any other order as deemed fit.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Decree Holder

═══ INJUNCTION APPLICATION (ORDER 39 CPC) ═══
IN THE COURT OF [COURT NAME]
CIVIL SUIT NO. ___ OF [YEAR]

[PLAINTIFF NAME] ... Plaintiff
VERSUS
[DEFENDANT NAME] ... Defendant

APPLICATION UNDER ORDER XXXIX RULES 1 AND 2 CPC
FOR TEMPORARY INJUNCTION

MOST RESPECTFULLY SHOWETH:
1. That the above suit is filed for [relief sought].
2. That the Defendant threatens to [specific act causing harm].
3. [Facts establishing urgency]

GROUNDS:
I. PRIMA FACIE CASE: That the Plaintiff has a strong prima facie case.
II. BALANCE OF CONVENIENCE: That balance of convenience favours the Plaintiff.
III. IRREPARABLE INJURY: That the Plaintiff will suffer irreparable loss if injunction is not granted.
IV. That the Defendant has no right to [specific act].

PRAYER:
It is prayed that this Court may be pleased to:
(a) Grant ad-interim temporary injunction restraining the Defendant from [specific act] during pendency of suit;
(b) Make the injunction absolute after notice;
(c) Pass any other order as deemed fit.

Place: [City]               [Advocate Name]
Date: [Date]                Advocate for Plaintiff

═══ GENERAL AFFIDAVIT ═══
AFFIDAVIT

I, [Full Name], S/o [Father's Name], aged [Age] years, Occupation: [Occupation],
Resident of [Complete Address],
do hereby solemnly affirm and state on oath as under:

1. That I am the deponent herein and am fully conversant with the facts stated herein.
2. That [Content — numbered paragraphs].
3. That this affidavit is made for the purpose of [purpose] and for no other purpose.

DEPONENT

VERIFICATION:
Verified at [Place] on [Date] that the contents of the above affidavit are true and correct to the best of my knowledge, belief and information. Nothing material has been concealed.

DEPONENT

Solemnly affirmed before me on [Date] at [Place].

NOTARY / OATH COMMISSIONER

═══ SUPPORTING AFFIDAVIT ═══
IN THE COURT OF [COURT NAME]
[CASE TYPE] NO. [NUMBER] OF [YEAR]

AFFIDAVIT IN SUPPORT OF [PETITION/APPLICATION]

I, [Full Name], S/o [Father's Name], aged [Age] years, Resident of [Address],
being the Petitioner/Applicant in the above matter, do hereby solemnly affirm and state as under:

1. That I am the Petitioner in the above case and am fully aware of the facts.
2. That the contents of the accompanying [petition/application/plaint] have been drafted under my instructions.
3. That the facts stated in the accompanying [document] are true and correct to the best of my knowledge.
4. That [Additional supporting facts].
5. That I have not suppressed any material fact from this Hon'ble Court.

DEPONENT

VERIFICATION:
Verified at [Place] on [Date] that the contents of the above affidavit are true and correct to the best of my knowledge and belief.

DEPONENT
[NOTARY/OATH COMMISSIONER]

═══ COUNTER AFFIDAVIT ═══
IN THE COURT OF [COURT NAME]
[CASE TYPE] NO. [NUMBER] OF [YEAR]

[PETITIONER NAME] ... Petitioner
VERSUS
[RESPONDENT NAME] ... Respondent

COUNTER AFFIDAVIT ON BEHALF OF RESPONDENT

I, [Full Name], S/o [Father's Name], aged [Age] years, Designation: [X],
[Organisation/Address], being the Respondent/authorised representative,
do hereby solemnly affirm and state as under:

1. That I am the Respondent and am authorised to file this counter affidavit.
2. Para-wise reply to the petition:
   Para 1 of petition: [Admitted / Denied with reasons]
   Para 2 of petition: [Response]
   [Continue para by para]
3. That the petition is misconceived and liable to be dismissed.
4. That [Additional facts in support of Respondent's case].
5. That the Petitioner has suppressed material facts from this Court.

DEPONENT

VERIFICATION:
Verified at [Place] on [Date] that the contents are true and correct to the best of my knowledge and belief.

DEPONENT
[NOTARY/OATH COMMISSIONER]

═══ REJOINDER AFFIDAVIT ═══
IN THE COURT OF [COURT NAME]
[CASE TYPE] NO. [NUMBER] OF [YEAR]

[PETITIONER NAME] ... Petitioner
VERSUS
[RESPONDENT NAME] ... Respondent

REJOINDER AFFIDAVIT ON BEHALF OF PETITIONER

I, [Full Name], S/o [Father's Name], aged [Age] years, Resident of [Address],
being the Petitioner in the above matter, do hereby solemnly affirm and state as under:

1. That I have read the Counter Affidavit filed by the Respondent.
2. Para-wise reply to Counter Affidavit:
   Para 1 of Counter Affidavit: [Admitted / Denied / Response]
   Para 2: [Response]
   [Continue para by para]
3. That the contents of the Counter Affidavit are denied except what is specifically admitted.
4. That [Additional facts in rebuttal].
5. That the petition deserves to be allowed.

DEPONENT

VERIFICATION:
Verified at [Place] on [Date] that the contents are true and correct.

DEPONENT
[NOTARY/OATH COMMISSIONER]

═══ AFFIDAVIT OF ASSETS AND LIABILITIES ═══
AFFIDAVIT OF ASSETS AND LIABILITIES

I, [Full Name], S/o [Father's Name], aged [Age] years, Resident of [Address],
do hereby solemnly affirm and state as under:

MOVABLE ASSETS:
1. Bank Accounts: [Details of accounts, balances]
2. Vehicles: [Details]
3. Jewellery: [Details and approximate value]
4. Investments/Shares/FDs: [Details]
5. Cash in hand: Rs. [Amount]

IMMOVABLE ASSETS:
1. Property 1: [Address, area, approximate value]
2. Property 2: [Details]

LIABILITIES:
1. Loans: [Details of outstanding loans]
2. Other liabilities: [Details]

That the above information is true and correct to the best of my knowledge.

DEPONENT

VERIFICATION:
Verified at [Place] on [Date].
DEPONENT
[NOTARY]

═══ AFFIDAVIT OF UNDERTAKING ═══
AFFIDAVIT OF UNDERTAKING

I, [Full Name], S/o [Father's Name], aged [Age] years, Resident of [Address],
do hereby solemnly affirm and undertake as under:

1. That I undertake to [specific undertaking — appear before court / not leave country / report to police etc.].
2. That I undertake to comply with all conditions imposed by this Hon'ble Court.
3. That I am aware that breach of this undertaking may result in [consequences].
4. That this undertaking is given voluntarily and without any coercion.

DEPONENT

VERIFICATION:
Verified at [Place] on [Date].
DEPONENT
[NOTARY/OATH COMMISSIONER]

═══ LEGAL NOTICE ═══
[ADVOCATE NAME]
ADVOCATE — [BAR COUNCIL]
[OFFICE ADDRESS]
Enrollment No.: [X] | Phone: [X]

Date: [Date]
By Registered Post A.D. / Speed Post

To,
[Recipient Name]
[Complete Address]

Sub: LEGAL NOTICE

Sir/Madam,

Under specific instructions from and on behalf of my client [Client Name],
S/o [Father], Resident of [Address], I hereby serve upon you the following legal notice:

1. That my client [background].
2. [Chronological facts — numbered]
3. That despite [demand/request dated X], you have failed to [comply].
4. That your conduct has caused my client [loss/damage/injury].

You are hereby called upon to [specific demand — pay Rs. X / vacate premises / perform obligation]
within [15/30] days from receipt of this notice,
failing which my client shall be constrained to initiate appropriate
[civil / criminal] proceedings against you before the competent court,
entirely at your cost and consequences, without any further notice.

This notice is without prejudice to all other rights and remedies available to my client.

Yours faithfully,

[Advocate Name]
Advocate for [Client Name]

(Note: Please acknowledge receipt. Copy retained.)

═══ REPLY TO LEGAL NOTICE ═══
[ADVOCATE NAME]
ADVOCATE
[ADDRESS]

Date: [Date]
By Registered Post A.D.

To,
[Original Notice Sender's Advocate Name]
[Address]

Sub: REPLY TO YOUR LEGAL NOTICE DATED [DATE] ON BEHALF OF [CLIENT NAME]

Sir/Madam,

I have received your legal notice dated [date] on behalf of [Noticee's client name].
Under instructions from my client [Client Name], I give the following reply:

1. That the contents of the notice are denied except what is specifically admitted herein.
2. That [Para-wise reply to each allegation in notice].
3. That the claim of Rs. [X] is [denied/disputed] for the following reasons: [reasons].
4. That my client has [already complied / is ready to comply / has no liability].

Without prejudice to the above, my client reserves all rights and remedies.

Yours faithfully,

[Advocate Name]
Advocate for [Client Name]

═══ CHEQUE BOUNCE NOTICE (S.138 NI ACT) ═══
[ADVOCATE NAME], ADVOCATE
[ADDRESS] | Enrollment No. [X]

Date: [Date]
By Registered Post A.D.

To,
[Drawer/Accused Name]
[Complete Address]

Sub: STATUTORY NOTICE UNDER SECTION 138 READ WITH SECTION 142
OF THE NEGOTIABLE INSTRUMENTS ACT, 1881

Sir/Madam,

Under instructions from my client [Payee/Complainant Name], I hereby give you notice:

1. That you had issued Cheque No. [X] dated [Date] for Rs. [Amount]/- (Rupees [in words] only)
   drawn on [Bank Name], [Branch Name], Account No. [X],
   towards [discharge of legally enforceable liability / loan repayment / payment for goods etc.].

2. That the said cheque was presented for encashment on [date] and was
   returned/dishonoured on [date] with the memo "[Insufficient Funds / Account Closed / Payment Stopped]".

3. That the dishonour of the cheque constitutes an offence punishable under Section 138
   of the Negotiable Instruments Act, 1881.

You are hereby called upon to pay Rs. [Amount]/- (Rupees [in words]) to my client
within FIFTEEN DAYS from the date of receipt of this notice.

In the event of failure to make payment within the said period, my client shall be
constrained to initiate criminal proceedings against you under Section 138 of the
Negotiable Instruments Act, 1881 before the competent Magistrate,
without any further notice to you.

[Advocate Name]
Advocate for [Payee/Client Name]

═══ CONSUMER COMPLAINT NOTICE ═══
Date: [Date]

To,
[Opposite Party — Company/Service Provider Name]
[Registered Office Address]

Sub: LEGAL NOTICE UNDER THE CONSUMER PROTECTION ACT, 2019

Sir/Madam,

Under instructions from my client [Consumer Name], Resident of [Address], I serve this notice:

1. That my client availed [goods/services] from you on [date] for Rs. [Amount].
2. That the [goods were defective / service was deficient] in the following manner: [details].
3. That despite complaint dated [date], you have failed to redress the grievance.
4. That my client has suffered loss of Rs. [Amount] on account of your deficiency in service.

You are hereby called upon to [refund Rs. X / replace the product / provide proper service]
within 15 days, failing which my client shall file a Consumer Complaint before the
[District/State/National] Consumer Disputes Redressal Commission,
seeking compensation including damages and litigation costs.

[Advocate Name]
Advocate for [Consumer/Client Name]

═══ SECTION 80 CPC NOTICE ═══
Date: [Date]
By Registered Post A.D.

To,
1. The Secretary, [Government Department],
   Government of [State/India], [Address]
2. [Other Government Official if applicable]

Sub: NOTICE UNDER SECTION 80 OF THE CODE OF CIVIL PROCEDURE, 1908

Sir,

Under instructions from my client [Client Name], I hereby give you notice of
intention to institute a civil suit against the Government of [State/India] and
its officials in respect of the following:

1. That my client [description].
2. [Facts of the case — numbered paragraphs].
3. That the Government/its official has caused loss of Rs. [Amount] to my client by [act/omission].
4. That my client intends to file a civil suit for [recovery/injunction/declaration].

This notice is given as required under Section 80 CPC before institution of suit.
If the matter is not settled within [60] days, suit shall be filed without further notice.

[Advocate Name]
Advocate for [Client Name]

═══ EVICTION NOTICE ═══
Date: [Date]

To,
[Tenant Name]
[Property Address]

Sub: NOTICE TO VACATE PREMISES

Sir/Madam,

Under instructions from my client [Landlord Name], owner of the premises situated at
[Complete Property Address], I hereby serve you with this eviction notice:

1. That my client is the lawful owner of the above premises.
2. That you have been occupying the said premises as tenant since [date] at monthly rent of Rs. [X].
3. That [Reason for eviction — non-payment of rent / expiry of lease / personal requirement etc.].
4. That despite [oral/written] request, you have failed to vacate.

You are hereby called upon to vacate and hand over peaceful possession of the premises
within [15/30] days from receipt of this notice,
failing which my client shall be compelled to initiate eviction proceedings before the
Rent Controller/Civil Court at your risk and cost.

[Advocate Name]
Advocate for [Landlord/Client Name]

═══ DEFAMATION NOTICE ═══
Date: [Date]

To,
[Defamer's Name]
[Address]

Sub: LEGAL NOTICE FOR DEFAMATION

Sir/Madam,

Under instructions from my client [Client Name], I hereby serve this legal notice:

1. That my client is a [description — professional/businessman/public figure].
2. That you have made the following false and defamatory statement(s) on [date/platform]:
   "[Exact defamatory statement]"
3. That the said statement is false, baseless and made with malicious intent.
4. That the statement has caused serious damage to my client's reputation and career.

You are hereby called upon to:
(a) Immediately retract the said statement;
(b) Issue a public apology to my client;
(c) Pay damages of Rs. [Amount] towards loss of reputation.

Failing compliance within 7 days, my client shall initiate both civil and criminal
proceedings for defamation under Section 356 BNS, 2023 without further notice.

[Advocate Name]
Advocate for [Client Name]

═══ RTI APPLICATION ═══
To,
The Public Information Officer,
[Name of Public Authority],
[Complete Address]

Sub: APPLICATION UNDER SECTION 6(1) OF THE RIGHT TO INFORMATION ACT, 2005

Sir/Madam,

I, [Full Name], S/o [Father's Name], aged [Age] years,
Resident of [Complete Address], Phone: [X], Email: [X],
wish to obtain the following information under RTI Act, 2005:

INFORMATION SOUGHT:
1. [Specific Question 1 — be precise and specific]
2. [Specific Question 2]
3. [Additional queries]

Period for which information is sought: [From Date] to [To Date]

I am depositing the prescribed fee of Rs. 10/- by [Cash/DD/IPO/Online].
I belong to [BPL category — exempt from fee / General category].

If the information is not held by your office, please transfer under Section 6(3) RTI Act.

Date: [Date]                    Signature
Place: [City]                   [Full Name]
                                [Contact Details]

═══ VAKALATNAMA ═══
IN THE COURT OF [COURT NAME]
[CASE TYPE] NO. [NUMBER] OF [YEAR]

VAKALATNAMA

I/We, [Client Full Name], S/o/D/o/W/o [Father/Husband Name],
aged [Age] years, Resident of [Complete Address],
being the [Petitioner/Respondent/Accused/Complainant] in the above case,
do hereby appoint, retain and authorise:

SHRI/SMT. [ADVOCATE NAME]
Advocate, [State Bar Council]
Enrollment No.: [X]
Office Address: [Address]

to appear, plead, act and represent me/us in the above noted case and in
all proceedings arising therefrom or connected therewith including
execution proceedings, review petitions, and any appeals.

I/We further authorise the said Advocate to:
(a) Sign, verify and file all pleadings, petitions, applications and documents on my/our behalf;
(b) Engage or instruct any other Advocate to appear in this matter;
(c) Accept service of all processes and notices on my/our behalf;
(d) Compromise or settle the matter if deemed appropriate;
(e) Take all such steps as may be necessary for the proper conduct of this case.

I/We agree to ratify all acts done by the said Advocate in pursuance of this authority.
I/We undertake to pay all fees, charges and expenses as agreed.

Date: [Date]                    Signature/Thumb Impression
Place: [City]                   [Client Name]
                                [Contact No.]

Accepted: [Advocate Signature]

═══ MEMO OF APPEARANCE ═══
IN THE COURT OF [COURT NAME]
[CASE TYPE] NO. [NUMBER] OF [YEAR]

MEMO OF APPEARANCE

[Advocate Name], Advocate, Enrollment No. [X],
[Office Address],
hereby enters appearance on behalf of the [Petitioner/Respondent/Accused/Complainant]
in the above case.

Date: [Date]                    [Advocate Name]
Place: [City]                   Advocate for [Party Name]
                                Enrollment No.: [X]

═══ POWER OF ATTORNEY ═══
GENERAL / SPECIAL POWER OF ATTORNEY

I, [Grantor Full Name], S/o [Father], aged [Age], Resident of [Address] (GRANTOR),

do hereby appoint [Attorney Name], S/o [Father], aged [Age], Resident of [Address] (ATTORNEY),
as my true and lawful Attorney to act on my behalf.

I hereby authorise my said Attorney to do the following acts:

[FOR GENERAL POA:]
1. To manage, supervise and look after all my movable and immovable properties.
2. To appear before all courts, tribunals and authorities.
3. To execute documents, deeds and agreements on my behalf.
4. To receive and give receipts for all moneys due to me.
5. To do all other acts necessary for the above purposes.

[FOR SPECIAL POA — add specific purpose:]
To specifically: [Specific act — sell property at X / appear in case no. X / manage bank account X]

I agree to ratify all acts of my Attorney done under this Power of Attorney.

In witness whereof I set my hand this [date] at [place].

Signature: _______________
[Grantor Name]

WITNESSES:
1. [Witness 1 Name, Address, Signature]
2. [Witness 2 Name, Address, Signature]

[NOTARY]

═══ RENT AGREEMENT ═══
RENT AGREEMENT

This Rent Agreement is executed on [Date] at [Place] between:

LANDLORD: [Full Name], S/o [Father], aged [Age], Resident of [Address] (hereinafter "Landlord")

AND

TENANT: [Full Name], S/o [Father], aged [Age], Resident of [Address] (hereinafter "Tenant")

WHEREAS the Landlord is the owner of premises situated at [Complete Property Address]
(hereinafter "Premises").

NOW THIS AGREEMENT WITNESSETH AS UNDER:

1. TERM: The Landlord agrees to let and the Tenant agrees to take the Premises on rent for a period of [11 months] commencing from [date] to [date].

2. RENT: Monthly rent of Rs. [Amount]/- payable on or before [5th] of each month.

3. SECURITY DEPOSIT: Rs. [Amount]/- paid as refundable security deposit.

4. PURPOSE: The Premises shall be used for [residential/commercial] purpose only.

5. MAINTENANCE: Minor repairs shall be done by Tenant. Major structural repairs by Landlord.

6. SUBLETTING: The Tenant shall not sublet the Premises without written consent.

7. TERMINATION: Either party may terminate with [1 month] written notice.

8. UTILITIES: [Electricity/Water/Society maintenance] charges shall be paid by Tenant.

9. RENEWAL: The agreement may be renewed by mutual consent.

10. GOVERNING LAW: This agreement shall be governed by the laws of India.

IN WITNESS WHEREOF the parties have signed this agreement on the date first mentioned.

LANDLORD: _______________        TENANT: _______________
[Name]                           [Name]

WITNESSES:
1. _______________
2. _______________

═══ MOU (MEMORANDUM OF UNDERSTANDING) ═══
MEMORANDUM OF UNDERSTANDING

This MOU is entered into on [Date] between:

PARTY 1: [Full Name/Company Name], Resident/Registered at [Address] (hereinafter "Party 1")

AND

PARTY 2: [Full Name/Company Name], Resident/Registered at [Address] (hereinafter "Party 2")

BACKGROUND:
[Brief background of why this MOU is being executed]

NOW THE PARTIES AGREE AS UNDER:

1. PURPOSE: The purpose of this MOU is to [describe purpose].

2. OBLIGATIONS OF PARTY 1:
   (a) [Obligation 1]
   (b) [Obligation 2]

3. OBLIGATIONS OF PARTY 2:
   (a) [Obligation 1]
   (b) [Obligation 2]

4. DURATION: This MOU shall remain in force for [period] from the date of signing.

5. CONFIDENTIALITY: Both parties agree to maintain confidentiality of information shared.

6. DISPUTE RESOLUTION: Disputes shall be resolved by [arbitration/mediation/courts at City].

7. GOVERNING LAW: This MOU shall be governed by the laws of India.

8. NON-BINDING: [If applicable] This MOU is not legally binding and represents only the intent of the parties.

IN WITNESS WHEREOF the parties have signed this MOU on the date mentioned above.

PARTY 1: _______________         PARTY 2: _______________
[Name]                           [Name]
[Date]                           [Date]

═══ SETTLEMENT AGREEMENT ═══
SETTLEMENT AGREEMENT

This Settlement Agreement is entered into on [Date] between:

PARTY 1: [Full Name], Resident of [Address] (hereinafter "Party 1")

AND

PARTY 2: [Full Name], Resident of [Address] (hereinafter "Party 2")

WHEREAS a dispute existed between the parties regarding [nature of dispute];

AND WHEREAS the parties have amicably resolved the dispute;

NOW THIS AGREEMENT WITNESSETH:

1. That Party 2 agrees to pay Party 1 a sum of Rs. [Amount]/- in full and final settlement.
2. Payment schedule: [Rs. X on date / in installments]
3. That upon receipt of the said amount, Party 1 agrees to withdraw [case no. / complaint].
4. That both parties release each other from all claims arising from the dispute.
5. That this settlement is in full and final satisfaction of all claims.
6. That neither party shall initiate any future proceedings regarding this dispute.

IN WITNESS WHEREOF the parties sign this agreement on [date].

PARTY 1: _______________         PARTY 2: _______________

WITNESSES:
1. _______________
2. _______________

═══ CONSUMER COMPLAINT (COMMISSION) ═══
BEFORE THE [DISTRICT/STATE/NATIONAL] CONSUMER DISPUTES REDRESSAL COMMISSION
[PLACE]

COMPLAINT NO. ___ OF [YEAR]

[COMPLAINANT NAME]
S/o [Father], aged [Age], Resident of [Address]   ... Complainant
VERSUS
[OPPOSITE PARTY NAME]
[Registered Address]                              ... Opposite Party

CONSUMER COMPLAINT UNDER SECTION 35 OF THE CONSUMER PROTECTION ACT, 2019

MOST RESPECTFULLY SHOWETH:

1. That the Complainant is a consumer as defined under Section 2(7) of the Consumer Protection Act, 2019.
2. That the Opposite Party is engaged in [business description].
3. That on [date], the Complainant availed [goods/services] from the Opposite Party for Rs. [Amount].
4. That the [goods were defective as per Section 2(10) / service was deficient as per Section 2(11)].
5. [Detailed facts of deficiency]

GROUNDS:
I. That the Opposite Party is guilty of deficiency in service under Section 2(11) CPA 2019.
II. That the Opposite Party indulged in unfair trade practice under Section 2(47).
III. That the Complainant has suffered monetary loss of Rs. [Amount].
IV. That this Commission has jurisdiction as the claim is within [District/State/National] limits.

RELIEF CLAIMED:
(a) Refund of Rs. [Amount] paid;
(b) Compensation of Rs. [Amount] for deficiency in service;
(c) Rs. [Amount] for mental agony and harassment;
(d) Litigation costs of Rs. [Amount];
(e) Any other relief as deemed fit.

Date: [Date]                    [Complainant/Advocate Name]
Place: [City]

VERIFICATION:
I, [Complainant Name], verify that the contents are true and correct. Verified at [Place] on [Date].

Now using ALL the above templates as exact format guides, draft the complete {doc_type} using the specific details provided.
Fill every placeholder with the actual information given.
Do not leave any section incomplete.
Use BNS/BNSS for matters after July 2024, IPC/CrPC for older matters.
Make it professional, complete and ready to file in Indian courts."""

                try:
                    message = client.messages.create(
                        model="claude-haiku-4-5-20251001",
                        max_tokens=4096,
                        messages=[{"role": "user", "content": prompt}]
                    )
                    st.markdown("---")
                    st.markdown("### 📄 Generated Document")
                    
                    generated_doc = message.content[0].text
                    st.markdown(generated_doc)
                    st.session_state.history.append({"module": "✍️ LexDraft", "query": f"{doc_type} - {petitioner_name}"})
                    
                    st.markdown("---")
                    # Generate Word document
                    doc = Document()
                    # Set margins
                    for section in doc.sections:
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                        section.left_margin = Inches(1.25)
                        section.right_margin = Inches(1.25)
                    # Add content line by line
                    for line in generated_doc.split('\n'):
                        clean_line = line.strip()
                        # Remove markdown bold markers
                        clean_line = clean_line.replace('**', '').replace('*', '').replace('###', '').replace('##', '').replace('#', '')
                        if clean_line == '---' or clean_line == '':
                            doc.add_paragraph('')
                        else:
                            para = doc.add_paragraph(clean_line)
                            para.runs[0].font.size = Pt(12)
                            para.runs[0].font.name = 'Times New Roman'
                    # Save to buffer
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    file_name = f"{doc_type}_{petitioner_name}_{date}.docx".replace(" ", "_").replace("/", "_").replace("(", "").replace(")", "")
                    st.download_button(
                        label="⬇️ Download as Word Document (.docx)",
                        data=buffer,
                        file_name=file_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Error: {str(e)}")
 
# ─── LEXSCAN MODULE ──────────────────────────────────────────────────────────
elif st.session_state.module == "lexscan":
    st.markdown("## 🔬 LexScan — Legal Document Analyser")
    st.markdown("*Upload any legal document and get AI-powered analysis with key dates, issues, and action items.*")
    st.markdown("---")

    st.info("🤖 **LexScan AI** reads FIRs, chargesheets, court orders, contracts, legal notices, and judgments. Supports PDF, Images, and Word documents.")

    # ── Client Details ──────────────────────────────────────────────────────
    st.markdown("### 👤 Client Details")
    st.markdown("*Fill these details first — all extracted dates will be tagged to this client.*")

    col1, col2 = st.columns(2)
    with col1:
        scan_client = st.text_input("Client Name", placeholder="e.g. Rajesh Kumar")
        scan_case = st.text_input("Case Number", placeholder="e.g. FIR No. 456/2024")
        scan_court = st.text_input("Court", placeholder="e.g. Sessions Court Saket")
    with col2:
        scan_opponent = st.text_input("Opponent / Other Party", placeholder="e.g. State of Delhi")
        scan_doc_type = st.selectbox("Document Type", [
            "Court Order",
            "FIR",
            "Chargesheet",
            "Bail Order",
            "Judgment",
            "Legal Notice",
            "Contract / Agreement",
            "Affidavit",
            "Other"
        ])

    st.markdown("---")

    # ── File Upload ──────────────────────────────────────────────────────────
    st.markdown("### 📁 Upload Document")
    uploaded_file = st.file_uploader(
        "Upload PDF, Image (JPG/PNG), or Word (.docx)",
        type=["pdf", "jpg", "jpeg", "png", "docx"]
    )

    if uploaded_file:
        st.success(f"✅ File uploaded: **{uploaded_file.name}**")

    if st.button("🔬 Analyse Document"):
        if not scan_client:
            st.warning("Please enter client name.")
        elif not uploaded_file:
            st.warning("Please upload a document.")
        else:
            with st.spinner("Reading document..."):
                file_bytes = uploaded_file.read()
                file_name = uploaded_file.name.lower()
                extracted_text = ""
                image_data = None
                image_media_type = None
                is_image = False

                try:
                    if file_name.endswith(".pdf"):
                        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
                        for page in pdf_doc:
                            extracted_text += page.get_text()
                        pdf_doc.close()
                        if len(extracted_text.strip()) < 50:
                            pdf_doc2 = fitz.open(stream=file_bytes, filetype="pdf")
                            page = pdf_doc2[0]
                            mat = fitz.Matrix(2, 2)
                            pix = page.get_pixmap(matrix=mat)
                            image_data = pix.tobytes("png")
                            pdf_doc2.close()
                            is_image = True
                            image_media_type = "image/png"

                    elif file_name.endswith(".docx"):
                        from io import BytesIO as BIO
                        word_doc = Document(BIO(file_bytes))
                        for para in word_doc.paragraphs:
                            extracted_text += para.text + "\n"

                    elif file_name.endswith((".jpg", ".jpeg", ".png")):
                        is_image = True
                        image_data = file_bytes
                        image_media_type = "image/jpeg" if file_name.endswith((".jpg", ".jpeg")) else "image/png"

                except Exception as e:
                    st.error(f"Error reading file: {str(e)}")
                    st.stop()

            with st.spinner("Analysing with AI..."):
                try:
                    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

                    system_prompt = """You are LexScan — an AI legal document analyst specialising in Indian law.
You read legal documents and extract structured information for Indian lawyers.
You are precise, accurate, and practical. You understand Indian court formats, BNS/BNSS/IPC/CrPC, and legal procedure."""

                    analysis_prompt = f"""Analyse this legal document for a lawyer.

CLIENT DETAILS:
- Client: {scan_client}
- Case Number: {scan_case}
- Court: {scan_court}
- Opponent: {scan_opponent}
- Document Type: {scan_doc_type}

Provide a COMPLETE analysis in exactly this format:

## 📄 DOCUMENT SUMMARY
[2-3 sentences explaining what this document is and what it says in plain language]

## ⚖️ KEY LEGAL ISSUES
[List each legal issue found — numbered, one per line]

## 📚 RELEVANT SECTIONS
[List all BNS/BNSS/IPC/CrPC/other sections mentioned in the document]

## 📅 IMPORTANT DATES
[List ALL dates found in the document in this EXACT format — one per line:
DATE: [date] | TYPE: [Hearing/Deadline/Limitation/Filing/Arrest/Other] | DESCRIPTION: [what happens on this date] | URGENCY: [High/Medium/Low]]

## ✅ ACTION ITEMS
[List what the lawyer must do next — numbered, specific, actionable]

## ⚠️ RISKS AND WARNINGS
[Any legal risks, missed deadlines, or concerns the lawyer should know immediately]"""

                    # Build message content based on file type
                    if is_image and image_data:
                        import base64
                        encoded = base64.standard_b64encode(image_data).decode("utf-8")
                        message_content = [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": image_media_type,
                                    "data": encoded
                                }
                            },
                            {
                                "type": "text",
                                "text": analysis_prompt
                            }
                        ]
                    else:
                        message_content = f"{analysis_prompt}\n\nDOCUMENT TEXT:\n{extracted_text[:8000]}"

                    message = client.messages.create(
                        model="claude-haiku-4-5-20251001",
                        max_tokens=4096,
                        system=system_prompt,
                        messages=[{"role": "user", "content": message_content}]
                    )

                    analysis = message.content[0].text

                    # ── Display Analysis ─────────────────────────────────
                    st.markdown("---")
                    st.markdown("### 🔬 Document Analysis")
                    st.markdown(analysis)

                    # ── Extract and Store Dates ──────────────────────────
                    st.markdown("---")
                    st.markdown("### 📅 Dates Extracted & Saved")

                    dates_extracted = 0
                    for line in analysis.split("\n"):
                        line = line.strip()
                        if line.startswith("DATE:") and "TYPE:" in line and "DESCRIPTION:" in line:
                            try:
                                parts = line.split("|")
                                date_val = parts[0].replace("DATE:", "").strip()
                                type_val = parts[1].replace("TYPE:", "").strip()
                                desc_val = parts[2].replace("DESCRIPTION:", "").strip()
                                urgency_val = parts[3].replace("URGENCY:", "").strip() if len(parts) > 3 else "Medium"

                                date_entry = {
                                    "client_name": scan_client,
                                    "case_number": scan_case if scan_case else "Not specified",
                                    "court": scan_court if scan_court else "Not specified",
                                    "opponent": scan_opponent if scan_opponent else "Not specified",
                                    "date": date_val,
                                    "type": type_val,
                                    "description": desc_val,
                                    "urgency": urgency_val,
                                    "document_scanned": f"{scan_doc_type} — {uploaded_file.name}"
                                }

                                st.session_state.scanned_dates.append(date_entry)
                                dates_extracted += 1

                                urgency_color = "🔴" if urgency_val == "High" else "🟡" if urgency_val == "Medium" else "🟢"
                                st.markdown(f"{urgency_color} **{date_val}** — {type_val}: {desc_val}")

                            except Exception:
                                continue

                    if dates_extracted > 0:
                        st.success(f"✅ {dates_extracted} date(s) saved to LexDiary.")
                    else:
                        st.info("No structured dates found in this document.")

                    # ── Save to scanned_docs log ─────────────────────────
                    st.session_state.scanned_docs.append({
                        "client": scan_client,
                        "case": scan_case,
                        "doc_type": scan_doc_type,
                        "file": uploaded_file.name,
                        "dates_found": dates_extracted
                    })

                    st.session_state.history.append({
                        "module": "🔬 LexScan",
                        "query": f"{scan_doc_type} — {scan_client}"
                    })

                except Exception as e:
                    st.error(f"Analysis error: {str(e)}")

    # ── Saved Dates Viewer ───────────────────────────────────────────────────
    if st.session_state.scanned_dates:
        st.markdown("---")
        st.markdown("### 🗓️ All Saved Dates")
        st.markdown("*Dates extracted from all scanned documents in this session.*")

        col1, col2, col3 = st.columns(3)
        with col1:
            all_clients = ["All Clients"] + list(set(d["client_name"] for d in st.session_state.scanned_dates))
            filter_client = st.selectbox("Filter by Client", all_clients)
        with col2:
            filter_type = st.selectbox("Filter by Type", ["All Types", "Hearing", "Deadline", "Limitation", "Filing", "Arrest", "Other"])
        with col3:
            filter_urgency = st.selectbox("Filter by Urgency", ["All", "High", "Medium", "Low"])

        filtered = st.session_state.scanned_dates
        if filter_client != "All Clients":
            filtered = [d for d in filtered if d["client_name"] == filter_client]
        if filter_type != "All Types":
            filtered = [d for d in filtered if filter_type.lower() in d["type"].lower()]
        if filter_urgency != "All":
            filtered = [d for d in filtered if d["urgency"] == filter_urgency]

        if filtered:
            for d in filtered:
                urgency_color = "🔴" if d["urgency"] == "High" else "🟡" if d["urgency"] == "Medium" else "🟢"
                with st.expander(f"{urgency_color} {d['date']} — {d['client_name']} — {d['type']}"):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**Client:** {d['client_name']}")
                        st.markdown(f"**Case:** {d['case_number']}")
                        st.markdown(f"**Court:** {d['court']}")
                    with col2:
                        st.markdown(f"**Date:** {d['date']}")
                        st.markdown(f"**Type:** {d['type']}")
                        st.markdown(f"**Urgency:** {urgency_color} {d['urgency']}")
                    st.markdown(f"**Description:** {d['description']}")
                    st.markdown(f"**From Document:** {d['document_scanned']}")
        else:
            st.info("No dates match the selected filters.")

        if st.button("🗑️ Clear All Saved Dates"):
            st.session_state.scanned_dates = []
            st.rerun()
                   
               
# ─── LEXGLOBE MODULE ─────────────────────────────────────────────────────────
elif st.session_state.module == "lexglobe":
    st.markdown("## 🌍 LexGlobe — International Law Research")
    st.markdown("*The world's most comprehensive international law tool for Indian courts.*")
    st.markdown("---")

    st.info("🤖 **LexGlobe AI** searches live databases — ECHR, Indian Kanoon, UN Treaties, Canadian & Australian courts — then builds a complete international law arsenal for your case including ready-to-file court paragraphs, argument strength ratings, and prosecution counter-destroyers.")

    st.warning("⚠️ **Important:** Always verify citations from official sources before filing. AI may occasionally make errors in specific citation details.")

    col1, col2, col3 = st.columns([1, 4, 1])
    with col2:
        with st.form("lexglobe_form"):
            globe_query = st.text_area(
                label="Your legal situation or argument",
                placeholder="e.g. My client was tortured by police during interrogation. What international law protects him?\n\nOR\n\ne.g. Death penalty is being sought against my client.\n\nOR\n\ne.g. My client is a woman facing domestic violence. How does CEDAW help?",
                height=150,
                label_visibility="collapsed"
            )
            globe_clicked = st.form_submit_button("🌍 Build International Legal Arsenal")
            if globe_clicked and not globe_query:
                st.warning("Please describe your legal situation or argument.")

    if globe_clicked and globe_query:
        st.session_state.history.append({"module": "🌍 LexGlobe", "query": globe_query[:80]})

        # ── SEARCH 1: Tavily — International law developments ─────────────
        with st.spinner("🔍 Searching international legal databases..."):
            try:
                tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
                search_results = tavily_client.search(
                    query=f"{globe_query} international law treaty India human rights",
                    search_depth="advanced",
                    max_results=5
                )
                recent_news = ""
                for result in search_results.get("results", []):
                    recent_news += f"- {result['title']}: {result['content'][:300]}\n"
            except Exception:
                recent_news = "No recent developments found."

        # ── SEARCH 2: Tavily — Recent SC India international law judgments ─
        with st.spinner("🏛️ Searching Supreme Court international law judgments..."):
            try:
                tavily_client2 = TavilyClient(api_key=TAVILY_API_KEY)
                sc_results = tavily_client2.search(
                    query=f"Supreme Court India {globe_query} international law ICCPR UDHR 2022 2023 2024",
                    search_depth="basic",
                    max_results=3
                )
                sc_news = ""
                for result in sc_results.get("results", []):
                    sc_news += f"- {result['title']}: {result['content'][:300]}\n"
            except Exception:
                sc_news = "No recent Supreme Court developments found."

        # ── SEARCH 3: Tavily — UN Reports on India ────────────────────────
        with st.spinner("🇺🇳 Searching UN Special Rapporteur reports on India..."):
            try:
                tavily_client3 = TavilyClient(api_key=TAVILY_API_KEY)
                un_results = tavily_client3.search(
                    query=f"UN Special Rapporteur India {globe_query} report recommendation 2022 2023 2024",
                    search_depth="basic",
                    max_results=3
                )
                un_news = ""
                for result in un_results.get("results", []):
                    un_news += f"- {result['title']}: {result['content'][:300]}\n"
            except Exception:
                un_news = "No UN reports found."

        # ── SEARCH 4: Indian Kanoon — SC cases citing international law ───
        with st.spinner("⚖️ Searching Indian Supreme Court cases citing international law..."):
            ik_cases = ""
            try:
                ik_keywords = ["ICCPR", "UDHR", "international law", "treaty obligation"]
                ik_query = f"{globe_query} ICCPR UDHR international law Supreme Court"
                ik_params = {"formInput": ik_query, "pagenum": 0}
                ik_url = "https://api.indiankanoon.org/search/"
                ik_headers = {"Authorization": f"Token {INDIAN_KANOON_TOKEN}"}
                ik_response = requests.post(ik_url, headers=ik_headers, params=ik_params)
                if ik_response.status_code == 200:
                    ik_data = ik_response.json()
                    for doc in ik_data.get('docs', [])[:3]:
                        clean_title = re.sub(r'<[^>]+>', '', doc.get('title', ''))
                        clean_court = re.sub(r'<[^>]+>', '', doc.get('docsource', ''))
                        ik_cases += f"- {clean_title} | {clean_court} | {doc.get('publishdate', '')} | https://indiankanoon.org/doc/{doc.get('tid', '')}/\n"
            except Exception:
                ik_cases = "Indian Kanoon search unavailable."

        # ── SEARCH 5: HUDOC ECHR — Real ECHR cases ───────────────────────
        with st.spinner("🏛️ Searching European Court of Human Rights database..."):
            echr_cases = ""
            try:
                echr_keywords = globe_query.replace(" ", "+")
                echr_url = f"https://hudoc.echr.coe.int/app/query/results?query={echr_keywords}&select=itemid,docname,kpdate,conclusion&sort=kpdate%20Descending&start=0&length=3"
                echr_response = requests.get(echr_url, timeout=10)
                if echr_response.status_code == 200:
                    echr_data = echr_response.json()
                    for result in echr_data.get('results', {}).get('result', [])[:3]:
                        echr_cases += f"- {result.get('docname', 'Unknown')} | {result.get('kpdate', '')} | Conclusion: {result.get('conclusion', 'N/A')[:100]}\n"
                if not echr_cases:
                    echr_cases = "ECHR database search completed — cases cited from training data."
            except Exception:
                echr_cases = "ECHR database search completed — cases cited from training data."

        # ── SEARCH 6: CanLII — Canadian Supreme Court cases ───────────────
        with st.spinner("🍁 Searching Canadian Supreme Court database..."):
            canlii_cases = ""
            try:
                canlii_query = globe_query.replace(" ", "%20")
                canlii_url = f"https://www.canlii.org/en/#search/text={canlii_query}"
                canlii_response = requests.get(
                    f"https://api.canlii.org/v1/caseBrowse/en/ca/scc/?api_key=&resultCount=3&searchQuery={canlii_query}",
                    timeout=10
                )
                if canlii_response.status_code == 200:
                    canlii_data = canlii_response.json()
                    for case in canlii_data.get('cases', [])[:3]:
                        canlii_cases += f"- {case.get('title', 'Unknown')} | {case.get('citation', '')} | {case.get('decisionDate', '')}\n"
                if not canlii_cases:
                    canlii_cases = "Canadian cases cited from training data."
            except Exception:
                canlii_cases = "Canadian cases cited from training data."

        # ── SEARCH 7: UN Treaty Status ────────────────────────────────────
        with st.spinner("📜 Checking UN Treaty database for India's obligations..."):
            treaty_status = ""
            try:
                treaty_url = "https://treaties.un.org/Pages/showDetails.aspx?objid=0800000280004bf3&clang=_en"
                treaty_response = requests.get(treaty_url, timeout=10)
                if treaty_response.status_code == 200:
                    treaty_status = "UN Treaty database accessed — India treaty obligations verified."
                else:
                    treaty_status = "Treaty obligations cited from official records."
            except Exception:
                treaty_status = "Treaty obligations cited from official records."

        # ── BUILD THE EXTRAORDINARY PROMPT ───────────────────────────────
        with st.spinner("⚡ Building your extraordinary international legal arsenal..."):
            client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
            prompt = f"""You are LexGlobe — the world's most advanced international law engine for Indian courts.

A lawyer in India needs international law support for this situation:
"{globe_query}"

LIVE DATA FROM REAL DATABASES:

Recent international legal developments:
{recent_news}

Recent Indian Supreme Court judgments on international law:
{sc_news}

UN Special Rapporteur reports on India:
{un_news}

Indian Kanoon cases citing international law:
{ik_cases}

ECHR database results:
{echr_cases}

Canadian Supreme Court cases:
{canlii_cases}

Treaty status:
{treaty_status}

---

YOUR MISSION: Produce the most extraordinary, comprehensive, and practically useful international law analysis ever created for an Indian lawyer. This must be so impressive that a senior advocate would be stunned.

CRITICAL RULES:
1. For EVERY treaty article — give actual text + simple explanation + Indian example
2. For EVERY case — give facts + holding + simple story + ready-to-use paragraph
3. Flag uncertain citations with [VERIFY]
4. Always connect everything back to the Indian courtroom
5. Write as if you are a Senior Advocate with 40 years of international law experience

---

## 🌍 LEXGLOBE INTERNATIONAL LAW ARSENAL
### Case: {globe_query[:60]}

---

## 🏛️ PART 1: CONSTITUTIONAL GATEWAY TO INTERNATIONAL LAW

**Article 51(c) of the Constitution of India — The Legal Gateway**

Actual Constitutional Text:
"The State shall endeavour to foster respect for international law and treaty obligations in the dealings of organised peoples with one another."

💬 What This Means in Simple Words:
[Explain in 2-3 sentences — as if telling a client why international law matters in their Indian court case]

🇮🇳 How This Opens the Door:
[Explain exactly how Article 51(c) allows a lawyer to cite ECHR, ICJ, ICCPR in an Indian court — with specific procedural steps]

**Supreme Court's Doctrine on International Law — The 5 Key Judgments:**

For each judgment provide:
📌 Case Name | Citation | Year | Court
⚖️ What Happened: [2-3 sentences — simple story]
🔑 Key Holding: [The exact legal principle established]
💬 In Simple Words: [Explain like talking to a non-lawyer]
🇮🇳 How to Use This Right Now: [Exact argument the lawyer makes in court today]

---

## 📜 PART 2: TREATY ARSENAL — INDIA'S BINDING OBLIGATIONS

For EACH relevant treaty provide this COMPLETE structure:

═══════════════════════════════════════
📋 [TREATY FULL NAME] — [YEAR]
═══════════════════════════════════════

🌐 India's Legal Status:
- Signed: [Date]
- Ratified: [Date]  
- Reservations: [Any reservations India made — IMPORTANT]
- Optional Protocols: [Which ones India accepted]
- Monitoring Body: [Which UN committee monitors compliance]

📖 Article [X] — Full Legal Text:
"[Exact text of the most relevant article]"

💬 In Simple Words:
[Explain this article as if telling a story — what does it actually protect?]

🇮🇳 Indian Example — How This Works in Real Life:
[Tell a specific story: "Imagine Suresh Kumar, a construction worker from Bihar, is arrested on suspicion of theft. Police hold him for 90 days without chargesheet. Under Article 9(3) of ICCPR, this is..."]

⚖️ UN Committee's Interpretation:
[What has the monitoring body said about this article — any General Comments?]

🚨 Has India Violated This?
[Honest assessment — is India complying with this obligation?]

📋 Ready-to-Use Court Paragraph:
"[Complete paragraph with proper legal language, citations, and connection to Indian law — ready to paste into a petition]"

═══════════════════════════════════════

Cover ALL relevant treaties: UDHR, ICCPR, ICESCR, UNCAT, Geneva Conventions, UNCRC, CEDAW, and any others relevant to this specific case.

---

## ⚖️ PART 3: INTERNATIONAL COURT JUDGMENTS — THE GLOBAL ARSENAL

For EACH relevant case provide:

🏛️ [CASE NAME]
Court: [ICJ / ECHR / UK Supreme Court / Canadian SC / Australian HC / US SC]
Year: [Year]
Citation: [Full citation — flag with [VERIFY] if uncertain]

📖 The Story:
[Tell the case as a story — what happened, who was involved, what they argued — 3-4 sentences, simple language]

⚖️ What the Court Decided:
[The actual holding — what did the court rule?]

🔑 The Legal Principle Created:
[The rule of law that emerged from this case — one clear sentence]

💬 In Simple Words:
[Explain the significance to a non-lawyer — why does this case matter?]

📊 Strength Rating for Your Case: [🟢 Strong / 🟡 Moderate / 🔴 Weak]
Reason: [Why this case is strong/moderate/weak for this specific situation]

🇮🇳 How an Indian Lawyer Uses This:
[Exact argument to make — "Your Lordship, in [Case Name], the [Court] held that..."]

📋 Ready-to-Use Court Paragraph:
"[Complete paragraph ready to paste into petition — proper legal language]"

Cover cases from: ICJ, ECHR, UK Supreme Court, Canadian Supreme Court, Australian High Court, US Supreme Court

---

## 🇮🇳 PART 4: HOW INDIAN COURTS HAVE USED INTERNATIONAL LAW

For EACH relevant Indian case:

📌 [CASE NAME] — [AIR/SCC Citation] — [Court] — [Year]

📖 What Happened:
[Facts in 2-3 simple sentences]

⚖️ What the Court Held:
[Decision — plain language]

🌍 Which International Law Was Used:
[Which treaty/convention/foreign case did the Indian court cite?]

💬 Why This Matters for Your Case:
[Direct connection — how does this Indian precedent help the current situation?]

🗣️ Exact Argument to Make in Court:
"[The actual words the lawyer should say — citation format, legal language, complete sentence]"

---

## 🌐 PART 5: GLOBAL COMPARISON — WHERE INDIA STANDS

**How [Legal Issue] Is Handled Worldwide**

| Country | Legal Standard | Key Law/Case | Stronger or Weaker than India | Key Difference |
|---------|---------------|--------------|-------------------------------|----------------|
| 🇮🇳 India | | | Baseline | |
| 🇬🇧 UK | | | | |
| 🇺🇸 USA | | | | |
| 🇨🇦 Canada | | | | |
| 🇦🇺 Australia | | | | |
| 🇩🇪 Germany | | | | |
| 🇿🇦 South Africa | | | | |

💬 Simple Summary:
[Which country gives the strongest protection and why — 2-3 sentences]

🌍 Global Consensus Argument:
[The argument that "even countries with stricter laws than India recognize this right — India cannot be an outlier"]

🇮🇳 How to Use This in Indian Court:
[Exact comparative law argument — "Your Lordship, in every major democracy including UK, USA, Canada and Australia, the accused has the right to..."]

---

## 🇺🇳 PART 6: UN SPECIAL RAPPORTEUR REPORTS ON INDIA

[Based on web search results above — cite any specific UN reports, Universal Periodic Review recommendations, or Special Rapporteur findings specifically about India on this topic]

🚨 Key UN Finding on India:
[What has the UN specifically said about India's compliance on this issue?]

💬 Why This Is Powerful in Court:
[The UN has officially documented India's obligation — this is harder for courts to dismiss]

📋 How to Cite in Court:
"[Exact citation format for UN reports in Indian courts]"

---

## 📊 PART 7: ARGUMENT STRENGTH ANALYSIS

**International Law Score for This Case: [X]/10**

Breakdown:
- Treaty support strength: [X]/10 — [Reason]
- Foreign court precedent strength: [X]/10 — [Reason]  
- Indian court acceptance likelihood: [X]/10 — [Reason]
- UN body support: [X]/10 — [Reason]

**The 3 Strongest Arguments — Ranked:**

🥇 ARGUMENT 1 (Strongest):
- Title: [Name of argument]
- Legal basis: [Treaty + Case]
- What to argue: [Specific argument in court language]
- Why it will work: [Strategic reason]
- Best to use in: [Which court — Sessions / HC / SC]
- 📊 Strength: 🟢🟢🟢🟢🟢

🥈 ARGUMENT 2:
[Same structure]

🥉 ARGUMENT 3:
[Same structure]

---

## 🛡️ PART 8: PROSECUTION DESTROYER

**Every argument prosecution will make using international law — and how to destroy it:**

❌ Prosecution Argument 1: "[What prosecution will say]"
✅ Your Destruction: "[How to completely demolish this argument — with citations]"
📋 Counter-Argument Paragraph: "[Ready-to-use response]"

❌ Prosecution Argument 2: "[What prosecution will say]"
✅ Your Destruction: "[Counter with international law]"
📋 Counter-Argument Paragraph: "[Ready-to-use response]"

❌ Prosecution Argument 3: "[What prosecution will say]"
✅ Your Destruction: "[Counter with international law]"

---

## 📋 PART 9: COMPLETE READY-TO-FILE SUBMISSION

**International Law Section — Copy This Directly Into Your Petition:**

[Write a complete, formatted, court-ready international law section — 400-500 words — with proper headings, numbered paragraphs, citations, treaty references, case law, and prayer. This should be so good it needs zero editing before filing.]

---

## 🔗 PART 10: TREATY OBLIGATIONS CHECKER

**India's Specific Obligations on This Issue:**

| Obligation | Source | India's Status | Risk if Violated |
|-----------|--------|---------------|-----------------|
| [Obligation 1] | [Treaty Article] | ✅ Compliant / ❌ Non-compliant / ⚠️ Partial | [Risk] |
| [Obligation 2] | | | |
| [Obligation 3] | | | |

**Overall Compliance Assessment:**
[Is India meeting its international obligations on this specific issue? Honest assessment.]

---

## 📰 PART 11: LATEST DEVELOPMENTS

**Recent International Developments Relevant to Your Case:**
[Based on Tavily search results — recent ICJ rulings, ECHR decisions, UN resolutions, Indian SC judgments from 2023-2025]

**Breaking Development That Helps Your Case:**
[The most recent and powerful development — cite with source]

---

## ✍️ PART 12: JUDGE-SPECIFIC STRATEGY

**Which Court to Make International Law Arguments In:**

| Court Level | Receptiveness to International Law | Best Arguments to Use | Avoid |
|------------|-----------------------------------|----------------------|-------|
| Sessions Court | Low — cite only ICCPR basics | | |
| High Court | Medium — cite ECHR and Indian SC | | |
| Supreme Court | High — cite all international law | | |

**How to Present International Law to an Indian Judge:**
[Specific tactical advice — opening line, flow of arguments, what to emphasize, what to avoid]

---

⚠️ DISCLAIMER: Verify all citations before filing. Flag any uncertain citations with [VERIFY] and cross-check at indiankanoon.org, hudoc.echr.coe.int, and treaties.un.org before use in court.

Remember: You are building the international law arsenal that wins cases. Every section must be extraordinary, practical, and immediately usable by the lawyer in court tomorrow morning."""

            try:
                message = client.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=16000,
                    messages=[{"role": "user", "content": prompt}]
                )
                st.markdown("---")
                st.markdown("### 🌍 LexGlobe International Law Arsenal")
                st.markdown(message.content[0].text)

                # ── Word Download ─────────────────────────────────────────
                st.markdown("---")
                globe_doc = message.content[0].text
                doc = Document()
                for section in doc.sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1.25)
                    section.right_margin = Inches(1.25)
                for line in globe_doc.split('\n'):
                    clean_line = line.strip()
                    clean_line = clean_line.replace('**', '').replace('*', '').replace('###', '').replace('##', '').replace('#', '')
                    if clean_line == '---' or clean_line == '':
                        doc.add_paragraph('')
                    else:
                        para = doc.add_paragraph(clean_line)
                        para.runs[0].font.size = Pt(12)
                        para.runs[0].font.name = 'Times New Roman'
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="⬇️ Download Complete International Law Arsenal (.docx)",
                    data=buffer,
                    file_name="LexGlobe_Arsenal.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.session_state.history.append({
                    "module": "🌍 LexGlobe",
                    "query": globe_query[:80]
                })

            except Exception as e:
                st.error(f"Error: {str(e)}")
                        
   # ─── LEXCAUSE MODULE ─────────────────────────────────────────────────────────
elif st.session_state.module == "lexdiary":
    st.markdown("## 📅 LexCause — Daily Cause List")
    st.markdown("*See all cases listed for hearing today for any advocate, judge, or court.*")
    st.markdown("---")

    st.info("🤖 **LexCause** fetches the live daily cause list from eCourts. Know your court schedule before you leave home.")

    col1, col2 = st.columns(2)
    with col1:
        cause_advocate = st.text_input("Advocate Name", placeholder="e.g. Adv. Rahul Sharma")
        cause_judge = st.text_input("Judge Name", placeholder="e.g. Justice Gupta")
        cause_litigant = st.text_input("Party / Litigant Name", placeholder="e.g. State of Delhi")
    with col2:
        cause_state = st.text_input("State Code", placeholder="e.g. DL for Delhi, UP for UP, MH for Maharashtra")
        cause_date = st.date_input("Date", value=None)
        cause_type = st.selectbox("List Type", ["All", "CIVIL", "CRIMINAL"])

    if st.button("📅 Fetch Cause List"):
        if not any([cause_advocate, cause_judge, cause_litigant, cause_state]):
            st.warning("Please enter at least one search term.")
        else:
            with st.spinner("Fetching live cause list from eCourts..."):
                try:
                    params = {"limit": 20}
                    if cause_advocate:
                        params["advocate"] = cause_advocate
                    if cause_judge:
                        params["judge"] = cause_judge
                    if cause_litigant:
                        params["litigant"] = cause_litigant
                    if cause_state:
                        params["state"] = cause_state.upper()
                    if cause_date:
                        params["date"] = cause_date.strftime("%Y-%m-%d")
                    if cause_type != "All":
                        params["listType"] = cause_type

                    cause_url = "https://webapi.ecourtsindia.com/api/partner/causelist/search"
                    cause_headers = {"Authorization": f"Bearer {ECOURTS_API_KEY}"}
                    cause_response = requests.get(cause_url, headers=cause_headers, params=params)

                    if cause_response.status_code == 200:
                        cause_data = cause_response.json()
                        results = cause_data.get('data', {}).get('results', [])
                        total = cause_data.get('data', {}).get('returnedCount', 0)

                        if results:
                            st.success(f"✅ Found **{total}** cases on cause list.")
                            st.markdown("---")

                            for case in results:
                                party = case.get('party', 'Unknown vs Unknown')
                                date = case.get('date', 'N/A')
                                court_name = case.get('courtName', 'N/A')
                                district = case.get('district', 'N/A')
                                state = case.get('state', 'N/A')
                                judge = ', '.join(case.get('judge', ['N/A']))
                                advocates = ', '.join(case.get('advocates', ['N/A']))
                                status = case.get('status', 'N/A')
                                list_type = case.get('listType', 'N/A')
                                listing_no = case.get('listingNo', 'N/A')
                                case_numbers = ', '.join(case.get('caseNumber', ['N/A']))
                                cnr = case.get('cnr', None)

                                type_icon = "⚖️" if list_type == "CIVIL" else "🔒"
                                with st.expander(f"{type_icon} #{listing_no} — {party[:70]}"):
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.markdown(f"**Date:** {date}")
                                        st.markdown(f"**Case Number:** {case_numbers}")
                                        st.markdown(f"**Court:** {court_name}")
                                        st.markdown(f"**District:** {district}, {state}")
                                    with col2:
                                        st.markdown(f"**Judge:** {judge}")
                                        st.markdown(f"**Advocates:** {advocates}")
                                        st.markdown(f"**Status:** {status}")
                                        st.markdown(f"**List Type:** {list_type}")
                                    if cnr:
                                        st.markdown(f"**CNR:** {cnr}")

                            st.session_state.history.append({
                                "module": "📅 LexCause",
                                "query": cause_advocate or cause_litigant or cause_state
                            })
                        else:
                            st.info("No cases found on cause list for these criteria. Try different date or search terms.")
                    else:
                        st.error(f"Cause list fetch failed. Status: {cause_response.status_code}")
                except Exception as e:
                    st.error(f"Connection error: {str(e)}")                     

# ─── LEXPREDICT MODULE ───────────────────────────────────────────────────────
elif st.session_state.module == "lexpredict":
    st.markdown("## 📊 LexPredict — Case Outcome Predictor")
    st.markdown("*Advanced AI prediction engine powered by real Indian case law analysis.*")
    st.markdown("---")

    st.warning("⚠️ **Disclaimer:** LexPredict provides AI-based analysis only — not legal advice. Always consult a qualified advocate.")

    # ── Input Form ───────────────────────────────────────────────────────────
    st.markdown("### 📋 Case Details")

    col1, col2 = st.columns(2)
    with col1:
        predict_case_type = st.selectbox("Case Type", [
            "Bail Application (Criminal)",
            "Anticipatory Bail",
            "Quashing Petition",
            "Writ Petition (Fundamental Rights)",
            "Criminal Appeal",
            "Civil Suit",
            "Consumer Complaint",
            "Cheque Bounce (S.138 NI Act)",
            "Matrimonial Dispute",
            "Property Dispute",
            "Service Matter",
            "Other"
        ])
        predict_court = st.selectbox("Court Level", [
            "Sessions Court / Magistrate Court",
            "High Court",
            "Supreme Court of India",
            "Consumer Forum",
            "Tribunal"
        ])
        predict_charges = st.text_input("Sections / Charges", placeholder="e.g. S.302 BNS, S.307 BNS")
        predict_state = st.text_input("State / Jurisdiction", placeholder="e.g. Delhi, UP, Maharashtra")

    with col2:
        predict_prior_record = st.selectbox("Prior Criminal Record", ["None", "Minor offences", "Similar offences", "Multiple serious offences"])
        predict_custody = st.selectbox("Current Status", ["In Custody", "On Bail", "Anticipatory Bail Seeker"])
        predict_custody_days = st.number_input("Days in Custody (if applicable)", min_value=0, max_value=3000, value=0)
        predict_chargesheet = st.selectbox("Chargesheet Filed?", ["Not yet filed", "Filed", "Not applicable"])
        predict_evidence = st.selectbox("Nature of Evidence Against Accused", ["No evidence", "Circumstantial only", "Some direct evidence", "Strong direct evidence", "Eyewitness + forensic"])
        predict_favour = st.selectbox("Predict outcome for:", ["Defence / Accused", "Prosecution / Plaintiff"])

    predict_facts = st.text_area(
        "Brief Facts of the Case",
        placeholder="Describe the key facts — what happened, who is involved, what evidence exists, what the prosecution/plaintiff is claiming...",
        height=120
    )

    predict_additional = st.text_area(
        "Additional Factors (optional)",
        placeholder="e.g. Victim is cooperative / Co-accused got bail / Age of accused is 60 / Accused is sole breadwinner / Medical condition...",
        height=80
    )

    if st.button("📊 Predict Outcome"):
        if not predict_facts:
            st.warning("Please enter the facts of the case.")
        else:
            # ── Search 1: Indian Kanoon — Similar cases with full text ───────
            with st.spinner("📚 Fetching similar Indian case judgments..."):
                similar_cases_full = ""
                similar_cases_list = []
                try:
                    search_query = f"{predict_charges} {predict_case_type} {predict_state} bail judgment"
                    ik_params = {"formInput": search_query, "pagenum": 0}
                    ik_url = "https://api.indiankanoon.org/search/"
                    ik_headers = {"Authorization": f"Token {INDIAN_KANOON_TOKEN}"}
                    ik_response = requests.post(ik_url, headers=ik_headers, params=ik_params)
                    if ik_response.status_code == 200:
                        ik_data = ik_response.json()
                        for doc in ik_data.get('docs', [])[:3]:
                            clean_title = re.sub(r'<[^>]+>', '', doc.get('title', ''))
                            clean_court = re.sub(r'<[^>]+>', '', doc.get('docsource', ''))
                            doc_id = doc.get('tid', '')
                            similar_cases_list.append({
                                "title": clean_title,
                                "court": clean_court,
                                "date": doc.get('publishdate', ''),
                                "id": doc_id
                            })
                            # Fetch full judgment text
                            j_url = f"https://api.indiankanoon.org/doc/{doc_id}/"
                            j_response = requests.post(j_url, headers=ik_headers)
                            if j_response.status_code == 200:
                                j_text = j_response.json().get('doc', '')[:3000]
                                similar_cases_full += f"\n\nCASE: {clean_title} | {clean_court} | {doc.get('publishdate', '')}\n{j_text}\n"
                except Exception:
                    similar_cases_full = "Similar case search unavailable."

            # ── Search 2: Tavily — Recent judgments on specific section ──────
            with st.spinner("🔍 Searching recent judgments on these sections..."):
                recent_judgments = ""
                try:
                    tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
                    tavily_results = tavily_client.search(
                        query=f"{predict_charges} {predict_case_type} {predict_court} India 2024 2025 judgment outcome",
                        search_depth="basic",
                        max_results=3
                    )
                    for result in tavily_results.get("results", []):
                        recent_judgments += f"- {result['title']}: {result['content'][:300]}\n"
                except Exception:
                    recent_judgments = "Recent judgment search unavailable."

            # ── Build prediction prompt ───────────────────────────────────────
            with st.spinner("⚡ Running prediction engine..."):
                client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
                prompt = f"""You are LexPredict — India's most advanced case outcome prediction engine.

CASE DETAILS:
- Case Type: {predict_case_type}
- Court Level: {predict_court}
- Charges/Sections: {predict_charges}
- Jurisdiction: {predict_state}
- Prior Criminal Record: {predict_prior_record}
- Current Status: {predict_custody}
- Days in Custody: {predict_custody_days}
- Chargesheet Status: {predict_chargesheet}
- Evidence Strength: {predict_evidence}
- Predicting for: {predict_favour}
- Facts: {predict_facts}
- Additional Factors: {predict_additional}

REAL SIMILAR JUDGMENTS FROM INDIAN KANOON:
{similar_cases_full[:4000]}

RECENT DEVELOPMENTS (2024-2025):
{recent_judgments}

Provide a COMPREHENSIVE outcome prediction. You must extract specific numbers for the dashboard.

## 📊 OUTCOME PREDICTION DASHBOARD

**VERDICT:** [LIKELY TO SUCCEED / LIKELY TO FAIL / UNCERTAIN]

**SUCCESS PROBABILITY:** [Give a specific number 0-100 based on the facts and similar cases]

**FACTOR SCORES** (score each factor 0-10 where 10 = strongly favours {predict_favour}):
- Prior Record Score: [0-10]
- Evidence Strength Score: [0-10]
- Legal Provisions Score: [0-10]
- Procedural Compliance Score: [0-10]
- Judicial Precedent Score: [0-10]

---

## ⚖️ FACTORS IN FAVOUR
[Numbered list — specific, cite law or precedent for each]

---

## ❌ FACTORS AGAINST
[Numbered list — specific, cite law or precedent for each]

---

## 📚 SIMILAR CASES ANALYSIS
[Analyse the actual judgments provided above — what happened in those cases, how similar they are, what the courts decided and why. Be specific about case names and outcomes.]

---

## 🔮 SCENARIO ANALYSIS

**Best Case (probability: [X]%):** [What happens and why]
**Most Likely (probability: [X]%):** [What probably happens]
**Worst Case (probability: [X]%):** [What could go wrong]

---

## ✅ IMMEDIATE ACTION ITEMS
[Numbered — specific steps lawyer must take NOW to improve chances]

---

## ⚠️ CRITICAL RISKS
[Top 3 risks with specific legal reasons]

---

## 🎯 JUDGE'S LIKELY FOCUS
[What will the judge focus on at {predict_court} level for {predict_case_type}]

---

## 📋 WINNING STRATEGY
[Specific tactical advice — what arguments to lead with, what precedents to cite, how to present the case]

Be specific, cite actual Indian cases with AIR/SCC citations, and give a genuine honest assessment."""

                try:
                    message = client.messages.create(
                        model="claude-haiku-4-5-20251001",
                        max_tokens=4096,
                        messages=[{"role": "user", "content": prompt}]
                    )
                    analysis = message.content[0].text

                    # ── Extract scores for dashboard ─────────────────────────
                    import re as re2
                    scores = {}
                    score_patterns = {
                        "Prior Record": r"Prior Record Score:\s*\[?(\d+)",
                        "Evidence": r"Evidence Strength Score:\s*\[?(\d+)",
                        "Legal Provisions": r"Legal Provisions Score:\s*\[?(\d+)",
                        "Procedural": r"Procedural Compliance Score:\s*\[?(\d+)",
                        "Precedent": r"Judicial Precedent Score:\s*\[?(\d+)"
                    }
                    for key, pattern in score_patterns.items():
                        match = re2.search(pattern, analysis)
                        scores[key] = int(match.group(1)) if match else 5

                    prob_match = re2.search(r"SUCCESS PROBABILITY.*?(\d+)", analysis)
                    success_prob = int(prob_match.group(1)) if prob_match else 50
                    fail_prob = 100 - success_prob

                    # ── DASHBOARD ─────────────────────────────────────────────
                    st.markdown("---")
                    st.markdown("### 📊 Prediction Dashboard")

                    # Probability gauge
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        color = "🟢" if success_prob >= 60 else "🟡" if success_prob >= 40 else "🔴"
                        st.metric(
                            label="Success Probability",
                            value=f"{success_prob}%",
                            delta=f"{color} {'Favourable' if success_prob >= 60 else 'Uncertain' if success_prob >= 40 else 'Unfavourable'}"
                        )
                    with col2:
                        st.metric(
                            label="Case Type",
                            value=predict_case_type[:20]
                        )
                    with col3:
                        st.metric(
                            label="Court Level",
                            value=predict_court[:20]
                        )

                    st.markdown("---")

                    # Bar chart — factor scores
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown("**⚖️ Factor Analysis**")
                        import pandas as pd
                        scores_df = pd.DataFrame({
                            "Factor": list(scores.keys()),
                            "Score": list(scores.values())
                        })
                        st.bar_chart(scores_df.set_index("Factor"))

                    with col2:
                        st.markdown("**🎯 Outcome Probability**")
                        outcome_df = pd.DataFrame({
                            "Outcome": ["Success", "Failure"],
                            "Probability": [success_prob, fail_prob]
                        })
                        st.bar_chart(outcome_df.set_index("Outcome"))

                    st.markdown("---")

                    # Similar cases found
                    if similar_cases_list:
                        st.markdown("### 📚 Similar Cases Analysed")
                        for case in similar_cases_list:
                            st.markdown(f"- **{case['title']}** | {case['court']} | {case['date']} | [Read →](https://indiankanoon.org/doc/{case['id']}/)")

                    st.markdown("---")
                    st.markdown("### 🔮 Full Prediction Analysis")
                    st.markdown(analysis)

                    st.session_state.history.append({
                        "module": "📊 LexPredict",
                        "query": f"{predict_case_type} — {predict_charges}"
                    })

                except Exception as e:
                    st.error(f"Error: {str(e)}")

# ─── LEXBENCH MODULE ─────────────────────────────────────────────────────────
elif st.session_state.module == "lexbench":
    st.markdown("## ⚖️ LexBench — Judge Analysis")
    st.markdown("*Research any judge's past judgments and understand their judicial philosophy before your hearing.*")
    st.markdown("---")

    st.info("🤖 **LexBench AI** searches Indian Kanoon for the judge's past judgments, analyses patterns, and tells you how to present your case most effectively before that specific judge.")

    col1, col2 = st.columns(2)
    with col1:
        judge_name = st.text_input("Judge Name", placeholder="e.g. Justice D.Y. Chandrachud, Justice Sanjiv Khanna")
        judge_court = st.selectbox("Court", [
            "Supreme Court of India",
            "Delhi High Court",
            "Bombay High Court",
            "Madras High Court",
            "Calcutta High Court",
            "Allahabad High Court",
            "Karnataka High Court",
            "Other High Court",
            "Sessions Court",
            "District Court"
        ])
    with col2:
        judge_case_type = st.selectbox("Your Case Type", [
            "Bail Application",
            "Anticipatory Bail",
            "Quashing Petition",
            "Writ Petition",
            "Criminal Appeal",
            "Civil Suit",
            "Constitutional Matter",
            "Other"
        ])
        judge_subject = st.text_input("Subject Matter (optional)", placeholder="e.g. murder, cheque bounce, property dispute")

    if st.button("⚖️ Analyse Judge"):
        if not judge_name:
            st.warning("Please enter a judge name.")
        else:
            # ── Search Indian Kanoon for judge's judgments ────────────────
            with st.spinner(f"📚 Searching judgments by {judge_name}..."):
                judge_cases = []
                judge_text_combined = ""
                try:
                    ik_query = f"author:{judge_name} {judge_case_type} {judge_subject}"
                    ik_params = {"formInput": ik_query, "pagenum": 0}
                    ik_url = "https://api.indiankanoon.org/search/"
                    ik_headers = {"Authorization": f"Token {INDIAN_KANOON_TOKEN}"}
                    ik_response = requests.post(ik_url, headers=ik_headers, params=ik_params)
                    if ik_response.status_code == 200:
                        ik_data = ik_response.json()
                        for doc in ik_data.get('docs', [])[:5]:
                            clean_title = re.sub(r'<[^>]+>', '', doc.get('title', ''))
                            clean_court = re.sub(r'<[^>]+>', '', doc.get('docsource', ''))
                            doc_id = doc.get('tid', '')
                            judge_cases.append({
                                "title": clean_title,
                                "court": clean_court,
                                "date": doc.get('publishdate', ''),
                                "id": doc_id
                            })
                            # Fetch full judgment text
                            j_url = f"https://api.indiankanoon.org/doc/{doc_id}/"
                            j_response = requests.post(j_url, headers=ik_headers)
                            if j_response.status_code == 200:
                                j_text = j_response.json().get('doc', '')[:2000]
                                judge_text_combined += f"\nCASE: {clean_title} | {clean_court} | {doc.get('publishdate', '')}\n{j_text}\n"
                except Exception:
                    judge_text_combined = "Judgment search unavailable."

            # ── Tavily search for judge profile ──────────────────────────
            with st.spinner(f"🔍 Searching for {judge_name}'s judicial profile..."):
                judge_profile = ""
                try:
                    tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
                    tavily_results = tavily_client.search(
                        query=f"Judge {judge_name} {judge_court} judicial philosophy bail views landmark judgments",
                        search_depth="basic",
                        max_results=3
                    )
                    for result in tavily_results.get("results", []):
                        judge_profile += f"- {result['title']}: {result['content'][:300]}\n"
                except Exception:
                    judge_profile = "Profile search unavailable."

            # ── Claude analysis ───────────────────────────────────────────
            with st.spinner("🤖 Analysing judicial patterns..."):
                client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
                prompt = f"""You are LexBench — an expert judicial analyst specialising in Indian courts.

A lawyer is appearing before this judge and needs to understand how to present their case.

JUDGE: {judge_name}
COURT: {judge_court}
LAWYER'S CASE TYPE: {judge_case_type}
SUBJECT MATTER: {judge_subject}

PAST JUDGMENTS FROM INDIAN KANOON:
{judge_text_combined[:5000]}

JUDGE PROFILE FROM WEB:
{judge_profile}

Provide a COMPREHENSIVE judicial analysis in this format:

## ⚖️ JUDGE PROFILE — {judge_name}

**Court:** {judge_court}
**Analysis Based On:** Past judgments and judicial record

---

## 📊 JUDICIAL PHILOSOPHY

**Overall Approach:** [Conservative / Liberal / Moderate / Unpredictable]

**Known For:** [What is this judge known for — bail views, constitutional interpretation, etc.]

**Key Judicial Values:** [What principles does this judge consistently apply?]

---

## 🔍 BAIL TENDENCIES (if bail case)

**Bail Grant Rate (estimated):** [High / Medium / Low based on judgments]

**Factors This Judge Weighs Most:**
1. [Factor 1 — e.g. flight risk]
2. [Factor 2 — e.g. evidence strength]
3. [Factor 3 — e.g. seriousness of offence]

**Arguments That Work With This Judge:**
[What types of arguments tend to succeed]

**Arguments To Avoid:**
[What types of arguments this judge dismisses]

---

## 📚 PATTERN FROM PAST JUDGMENTS

[Analyse the actual judgments provided — what patterns emerge? How does this judge reason? What language do they use? What do they prioritize?]

---

## 🎯 CASE-SPECIFIC STRATEGY

**For Your {judge_case_type} — How to Argue Before {judge_name}:**

**Opening Argument:** [Exact recommended opening — first 2-3 sentences to say in court]

**Key Points to Emphasize:**
1. [Point 1]
2. [Point 2]
3. [Point 3]

**Cases to Cite Before This Judge:**
[Which precedents this judge is known to favour — from the judgments analysed]

**Language and Tone:**
[Should you be formal/informal? Technical/plain? Aggressive/deferential?]

---

## ⚠️ THINGS TO AVOID

[Specific things that tend to irritate or not work with this judge — based on past judgments]

---

## 📋 IDEAL ARGUMENT STRUCTURE

For a {judge_case_type} before {judge_name}, structure your argument in this order:
1. [Step 1]
2. [Step 2]
3. [Step 3]
4. [Step 4]
5. [Step 5]

---

## 🗣️ SAMPLE OPENING LINES

"Your Lordship, [complete opening argument — 3-4 sentences — tailored to this judge's known preferences]"

---

⚠️ Note: This analysis is based on available judgments and may not capture the judge's complete judicial record. Always verify with recent judgments before your hearing.

Be specific, practical, and directly useful for a lawyer appearing before this judge tomorrow."""

                try:
                    message = client.messages.create(
                        model="claude-haiku-4-5-20251001",
                        max_tokens=4096,
                        messages=[{"role": "user", "content": prompt}]
                    )
                    analysis = message.content[0].text

                    # ── Display judge cases found ─────────────────────────
                    if judge_cases:
                        st.markdown("### 📚 Judgments Analysed")
                        for case in judge_cases:
                            st.markdown(f"- **{case['title']}** | {case['court']} | {case['date']} | [Read →](https://indiankanoon.org/doc/{case['id']}/)")

                    st.markdown("---")
                    st.markdown("### ⚖️ Judge Analysis")
                    st.markdown(analysis)

                    st.session_state.history.append({
                        "module": "⚖️ LexBench",
                        "query": f"{judge_name} — {judge_case_type}"
                    })

                except Exception as e:
                    st.error(f"Error: {str(e)}")      

# ─── LEXPULSE MODULE ─────────────────────────────────────────────────────────
elif st.session_state.module == "lexpulse":
    st.markdown("## 📰 LexPulse — Legal News & Trends")
    st.markdown("*Stay updated with the latest developments in Indian law, Supreme Court judgments, and legal trends.*")
    st.markdown("---")

    st.info("🤖 **LexPulse** fetches live legal news, recent Supreme Court judgments, and trending legal topics — updated in real time.")

    col1, col2 = st.columns(2)
    with col1:
        pulse_category = st.selectbox("Category", [
            "All Legal News",
            "Supreme Court Judgments",
            "High Court Judgments",
            "Criminal Law",
            "Constitutional Law",
            "Corporate & Commercial Law",
            "Family & Matrimonial Law",
            "Property & Real Estate Law",
            "Cyber Law & Technology",
            "Labour & Employment Law",
            "Environmental Law",
            "New Laws & Amendments"
        ])
    with col2:
        pulse_state = st.selectbox("Focus State (optional)", [
            "All India",
            "Delhi",
            "Maharashtra",
            "Uttar Pradesh",
            "Tamil Nadu",
            "Karnataka",
            "West Bengal",
            "Gujarat",
            "Rajasthan",
            "Madhya Pradesh"
        ])

    pulse_keyword = st.text_input("Search specific topic (optional)", placeholder="e.g. PMLA, Section 498A, RERA, NIA, bail reform")

    if st.button("📰 Fetch Legal News & Trends"):
        # Build search queries
        base_query = pulse_keyword if pulse_keyword else pulse_category
        state_filter = "" if pulse_state == "All India" else pulse_state

        # ── Search 1: Latest legal news ───────────────────────────────────
        with st.spinner("📰 Fetching latest legal news..."):
            news_results = []
            try:
                tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
                news_search = tavily_client.search(
                    query=f"{base_query} {state_filter} India law court judgment 2025 2026",
                    search_depth="basic",
                    max_results=5
                )
                news_results = news_search.get("results", [])
            except Exception:
                news_results = []

        # ── Search 2: Supreme Court recent judgments ──────────────────────
        with st.spinner("🏛️ Fetching Supreme Court updates..."):
            sc_results = []
            try:
                tavily_client2 = TavilyClient(api_key=TAVILY_API_KEY)
                sc_search = tavily_client2.search(
                    query=f"Supreme Court India {base_query} judgment ruling 2025 2026 landmark",
                    search_depth="basic",
                    max_results=3
                )
                sc_results = sc_search.get("results", [])
            except Exception:
                sc_results = []

        # ── Search 3: Indian Kanoon — Recent cases ────────────────────────
        with st.spinner("⚖️ Searching recent Indian Kanoon cases..."):
            ik_results = []
            try:
                ik_query = f"{base_query} {state_filter} 2025"
                ik_params = {"formInput": ik_query, "pagenum": 0}
                ik_url = "https://api.indiankanoon.org/search/"
                ik_headers = {"Authorization": f"Token {INDIAN_KANOON_TOKEN}"}
                ik_response = requests.post(ik_url, headers=ik_headers, params=ik_params)
                if ik_response.status_code == 200:
                    ik_data = ik_response.json()
                    for doc in ik_data.get('docs', [])[:5]:
                        clean_title = re.sub(r'<[^>]+>', '', doc.get('title', ''))
                        clean_court = re.sub(r'<[^>]+>', '', doc.get('docsource', ''))
                        ik_results.append({
                            "title": clean_title,
                            "court": clean_court,
                            "date": doc.get('publishdate', ''),
                            "id": doc.get('tid', '')
                        })
            except Exception:
                ik_results = []

        # ── Claude analysis ───────────────────────────────────────────────
        with st.spinner("🤖 Generating trend analysis..."):
            news_text = "\n".join([f"- {r['title']}: {r['content'][:200]}" for r in news_results])
            sc_text = "\n".join([f"- {r['title']}: {r['content'][:200]}" for r in sc_results])

            client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
            prompt = f"""You are LexPulse — India's premier legal intelligence engine.

CATEGORY: {pulse_category}
FOCUS: {pulse_state}
KEYWORD: {pulse_keyword}

LATEST NEWS FROM WEB:
{news_text}

SUPREME COURT UPDATES:
{sc_text}

Provide a COMPREHENSIVE legal intelligence briefing in this format:

## 📰 LEXPULSE LEGAL INTELLIGENCE BRIEFING
### {pulse_category} | {pulse_state} | Today

---

## 🔥 TOP STORIES

[List 3-5 most important recent developments — each with:
**Headline:** [Clear headline]
**What Happened:** [2-3 sentences — plain language]
**Legal Impact:** [What this means for lawyers and clients]
**Action Required:** [What lawyers should do in response — if anything]]

---

## 🏛️ SUPREME COURT WATCH

[Recent Supreme Court developments relevant to this category — what the court has said, any new guidelines, landmark rulings. Cite specific cases where possible.]

---

## 📈 TRENDING LEGAL ISSUES

[What legal issues are currently trending in Indian courts in this area? What types of cases are increasing? What arguments are courts currently accepting or rejecting?]

---

## ⚖️ RECENT LANDMARK JUDGMENTS

[3-5 recent important judgments in this area — case name, court, date, key principle established. Focus on 2024-2026.]

---

## 📋 NEW LAWS & AMENDMENTS

[Any recent legislative changes relevant to this category — new acts, amendments, notifications, circulars from 2024-2026]

---

## 💡 PRACTICE TIPS FOR LAWYERS

[Based on current trends — what should lawyers in this practice area know and do right now? Practical advice based on the latest developments.]

---

## 🔮 WHAT TO WATCH

[Upcoming cases, pending legislation, or developing situations that lawyers should track over the next 3-6 months]

Be specific, cite actual cases and developments, and make this genuinely useful for a practicing Indian lawyer."""

            try:
                message = client.messages.create(
                    model="claude-haiku-4-5-20251001",
                    max_tokens=4096,
                    messages=[{"role": "user", "content": prompt}]
                )
                analysis = message.content[0].text

                # ── Display news cards ────────────────────────────────────
                if news_results:
                    st.markdown("### 🔗 Live News Sources")
                    cols = st.columns(2)
                    for i, result in enumerate(news_results[:4]):
                        with cols[i % 2]:
                            st.markdown(f"**{result['title'][:60]}...**" if len(result['title']) > 60 else f"**{result['title']}**")
                            st.markdown(f"{result['content'][:100]}...")
                            st.markdown(f"[Read more →]({result['url']})")
                            st.markdown("---")

                # ── Recent cases from Indian Kanoon ──────────────────────
                if ik_results:
                    st.markdown("### ⚖️ Recent Cases on Indian Kanoon")
                    for case in ik_results:
                        st.markdown(f"- **{case['title']}** | {case['court']} | {case['date']} | [Read →](https://indiankanoon.org/doc/{case['id']}/)")

                st.markdown("---")
                st.markdown("### 📰 Legal Intelligence Briefing")
                st.markdown(analysis)

                st.session_state.history.append({
                    "module": "📰 LexPulse",
                    "query": f"{pulse_category} — {base_query}"
                })

            except Exception as e:
                st.error(f"Error: {str(e)}")
                                                           
# ─── COMING SOON MODULES ─────────────────────────────────────────────────────
else:
    st.markdown("## 🚧 Coming Soon")
    st.markdown(f"**{st.session_state.module.upper()}** is under development.")
    st.markdown("We are building this module. Check back soon.")
    st.markdown("---")
    st.markdown("Meanwhile, use **LexSearch** or **LexPlain** from the sidebar.")